"""Build the auditable weekly-planning source and quality register."""

from __future__ import annotations

import csv
import json
import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[3]
DATA = ROOT / "data"
OUTPUT = ROOT / "output"
SCHEMA_PATH = DATA / "schemas" / "review_register_schema.json"


def catalog_rows(fields: list[str]) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for path in sorted((DATA / "data" / "catalog").glob("annual_record_bundle_*.json")):
        payload = json.loads(path.read_text())
        for source_row, record in enumerate(payload["records"], start=1):
            year = int(record["refYear"])
            flow = record["flowCode"]
            is_total = (
                record.get("partnerCode") == 0
                and record.get("partner2Code") == 0
                and record.get("motCode") == 0
            )
            coverage_note = "Official catalog record"
            quality_id = ""
            if year == 2019:
                coverage_note = "2019 catalog bundle is import-only; no official export records"
                quality_id = "QUAL-COMTRADE-GAP-01"
            row = {
                "record_id": f"TRD-{year}-{flow}-{source_row:04d}",
                "source_release": path.name,
                "source_row": source_row,
                "source_classification": "official_catalog",
                "source_record_stage": "pre_analysis_catalog_record",
                "source_coverage_note": coverage_note,
                "ref_year": year,
                "flow_code": flow,
                "partner_code": record.get("partnerCode"),
                "partner2_code": record.get("partner2Code"),
                "mot_code": record.get("motCode"),
                "primary_value": record.get("primaryValue"),
                "is_reported": record.get("isReported"),
                "is_aggregate": record.get("isAggregate"),
                "legacy_estimation_flag": record.get("legacyEstimationFlag"),
                "inclusion_status": (
                    "included_headline" if is_total else "enumerated_not_headline"
                ),
                "quality_decision_id": quality_id,
                "governing_rule": (
                    "REV-MR-22-02 section 2; retained by REV-MR-23-03 section 3; all-world total row"
                    if is_total
                    else "REV-MR-22-02 section 2; retained by REV-MR-23-03 section 3; enumerate before analysis; no route-row stacking"
                ),
            }
            rows.append({field: row.get(field, "") for field in fields})
    return rows


def bridge_rows(fields: list[str]) -> list[dict[str, object]]:
    path = DATA / "data" / "trade" / "governed_export_bridge_2018_2022.csv"
    rows: list[dict[str, object]] = []
    with path.open(newline="") as handle:
        for source_row, bridge in enumerate(csv.DictReader(handle), start=2):
            year = int(bridge["year"])
            note = (
                "2019 catalog bundle is import-only; export is a governed bridge"
                if year == 2019
                else f"No official {year} export catalog record; governed export bridge"
            )
            row = {
                "record_id": bridge["bridge_obs_id"],
                "source_release": path.name,
                "source_row": source_row,
                "source_classification": "modeled_export_bridge",
                "source_record_stage": "governed_model_output",
                "source_coverage_note": note,
                "ref_year": year,
                "flow_code": "X",
                "partner_code": 0,
                "partner2_code": 0,
                "mot_code": 0,
                "primary_value": bridge["export_value_usd"],
                "is_reported": False,
                "is_aggregate": True,
                "legacy_estimation_flag": "",
                "inclusion_status": "included_bridge",
                "quality_decision_id": "QUAL-COMTRADE-GAP-01",
                "governing_rule": "REV-MR-23-03; MODEL-BRIDGE-23-03",
            }
            rows.append({field: row.get(field, "") for field in fields})
    return rows


def quality_decisions(fields: list[str]) -> list[dict[str, object]]:
    decisions = [
        {
            "quality_decision_id": "QD-CENSUS-HTML-EXCLUDE",
            "evidence_id": "QUAL-CENSUS-HTML-01",
            "source_scope": "supply_chain/census_trade/",
            "decision": "exclude_invalid_html",
            "inclusion_status": "excluded",
            "affected_records": 1,
            "rationale": "Capture is an invalid-key HTML response and is not trade evidence.",
            "governing_rule": "REV-MR-23-03 source quality",
        },
        {
            "quality_decision_id": "QD-BLS-EMPTY-EXCLUDE",
            "evidence_id": "QUAL-BLS-EMPTY-01",
            "source_scope": "supply_chain/bls/",
            "decision": "exclude_empty_series",
            "inclusion_status": "excluded",
            "affected_records": 0,
            "rationale": "BLS stub contains no observations; populated ILO labor records remain eligible.",
            "governing_rule": "REV-MR-23-03 source quality",
        },
        {
            "quality_decision_id": "QD-COMTRADE-GAP-BRIDGE",
            "evidence_id": "QUAL-COMTRADE-GAP-01",
            "source_scope": "supply_chain/comtrade/ GBR HS29 exports",
            "decision": "enumerate_official_records_then_bridge_missing_exports",
            "inclusion_status": "conditional_bridge",
            "affected_records": 5,
            "rationale": "Official export coverage is absent for 2018-2022; the 2019 bundle contains imports only.",
            "governing_rule": "REV-MR-23-03; MODEL-BRIDGE-23-03",
        },
    ]
    return [{field: row.get(field, "") for field in fields} for row in decisions]


def governance_rows(fields: list[str]) -> list[dict[str, object]]:
    rows = [
        {
            "parameter_id": "headline_trade_aggregation",
            "current_governing_version": "REV-MR-22-02 section 2 (retained by REV-MR-23-03 section 3)",
            "current_value": "Use the all-world total row for each year and flow when it exists.",
            "superseded_version": "REV-MR-22-01 section 2",
            "superseded_value": "All route rows may be counted.",
            "result_changed_by_supersession": True,
            "result_effect": "Prevents route-row stacking from inflating headline totals.",
            "governing_evidence": "REV-MR-22-02; REV-MR-23-03 supersession notes; 2022-11-21 method-release email",
        },
        {
            "parameter_id": "partner_aggregation",
            "current_governing_version": "REV-MR-22-02 section 2 (retained by REV-MR-23-03 section 3)",
            "current_value": "Do not stack transport-mode rows into a country total.",
            "superseded_version": "REV-MR-22-01 section 2",
            "superseded_value": "All route rows may be counted as country observations.",
            "result_changed_by_supersession": True,
            "result_effect": "Changes partner totals and rankings where multiple transport modes exist.",
            "governing_evidence": "REV-MR-22-02; REV-MR-23-03 supersession notes; 2022-11-21 method-release email",
        },
        {
            "parameter_id": "forecast_basis",
            "current_governing_version": "REV-MR-23-03 section 2",
            "current_value": "Bridge missing exports with price and macro context; forecast from eight-year bridged real-volume history.",
            "superseded_version": "REV-MR-22-02 forecast basis",
            "superseded_value": "Macro bridge only when official coverage is absent; earlier v1 used a two-year compound annual rate.",
            "result_changed_by_supersession": True,
            "result_effect": "Changes bridge and forecast values, including the modeled 2018-2022 export series.",
            "governing_evidence": "REV-MR-23-03 section 2 and section 3; MODEL-BRIDGE-23-03",
        },
        {
            "parameter_id": "review_threshold_gbp",
            "current_governing_version": "REV-MR-23-03 section 3",
            "current_value": "35000 GBP risk-adjusted exposure at ERP-code level.",
            "superseded_version": "REV-MR-22-02 review threshold",
            "superseded_value": "30000 GBP.",
            "result_changed_by_supersession": True,
            "result_effect": "Exposures from 30000 through 34999.99 GBP no longer cross the review threshold.",
            "governing_evidence": "REV-MR-23-03 section 3; material_catalog.json",
        },
        {
            "parameter_id": "source_quality",
            "current_governing_version": "REV-MR-23-03 section 2",
            "current_value": "Exclude invalid Census captures and empty BLS stubs.",
            "superseded_version": "",
            "superseded_value": "No explicit prior parameter value.",
            "result_changed_by_supersession": False,
            "result_effect": "New explicit rule; removes unsupported artifacts from evidence rather than changing a prior governed result.",
            "governing_evidence": "REV-MR-23-03; QUAL-CENSUS-HTML-01; QUAL-BLS-EMPTY-01",
        },
        {
            "parameter_id": "coverage_gate",
            "current_governing_version": "REV-MR-23-03 section 2",
            "current_value": "Show record population before ranking partners.",
            "superseded_version": "",
            "superseded_value": "Enumeration was already required in REV-MR-22-01 and REV-MR-22-02.",
            "result_changed_by_supersession": False,
            "result_effect": "No result change; v3 restates the existing enumeration requirement as a coverage gate.",
            "governing_evidence": "REV-MR-23-03; REV-MR-22-02; REV-MR-22-01",
        },
        {
            "parameter_id": "recovery_benchmark",
            "current_governing_version": "REV-MR-23-03 section 2",
            "current_value": "92 percent of pre-shock bridge baseline.",
            "superseded_version": "",
            "superseded_value": "No prior benchmark.",
            "result_changed_by_supersession": False,
            "result_effect": "New parameter; establishes a recovery classification but does not supersede a prior result.",
            "governing_evidence": "REV-MR-23-03 section 2 and section 3",
        },
    ]
    return [{field: row.get(field, "") for field in fields} for row in rows]


def entity_resolution_rows(fields: list[str]) -> list[dict[str, object]]:
    appendix = Document(DATA / "docs" / "procurement" / "supplier_catalog_appendix.docx")
    table = appendix.tables[1]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    appendix_rows = [
        dict(zip(headers, [cell.text.strip() for cell in row.cells]))
        for row in table.rows[1:]
    ]

    materials_path = DATA / "config" / "procurement" / "material_catalog.json"
    materials = {
        row["erp_code"]: row
        for row in json.loads(materials_path.read_text())["records"]
    }

    orders_path = DATA / "data" / "procurement" / "weekly_order_board.xlsx"
    order_book = load_workbook(orders_path, read_only=True, data_only=True)
    order_sheet = order_book["Orders"]
    order_headers = [cell.value for cell in next(order_sheet.iter_rows(min_row=1, max_row=1))]
    orders: dict[str, list[dict[str, object]]] = {}
    for values in order_sheet.iter_rows(min_row=2, values_only=True):
        order = dict(zip(order_headers, values))
        orders.setdefault(str(order["erp_code"]), []).append(order)

    resolved: list[dict[str, object]] = []
    for appendix_row in appendix_rows:
        erp_code = appendix_row["ERP code"]
        material = materials[erp_code]
        entity_orders = orders.get(erp_code, [])
        assert material["bom_ref"] == appendix_row["BOM reference"]
        assert material["supplier_portal_id"] == appendix_row["Supplier portal ID"]
        assert material["hs_tariff_code"] == appendix_row["HS tariff code"]
        assert all(order["portal_id"] == appendix_row["Supplier portal ID"] for order in entity_orders)
        assert all(order["supplier_name"] == appendix_row["Tier-1 supplier"] for order in entity_orders)
        row = {
            "canonical_entity_id": f"MAT-{erp_code.removeprefix('ERP-OC-')}",
            "erp_code": erp_code,
            "procurement_po_numbers": "; ".join(str(order["po_number"]) for order in entity_orders),
            "supplier_portal_id": appendix_row["Supplier portal ID"],
            "bom_reference": appendix_row["BOM reference"],
            "hs_tariff_code": appendix_row["HS tariff code"],
            "hs_chapter": "29",
            "tier1_supplier": appendix_row["Tier-1 supplier"],
            "trade_partner_market_code": appendix_row["Partner market code"],
            "trade_scope_id": "CAT-GBR-HS29-REFRESH-2023-02",
            "mapping_status": "resolved",
            "resolution_note": "Canonical material/supplier crosswalk. Trade observations are HS29 chapter-level; partner market and tariff identifiers must both be retained and must not be treated as interchangeable.",
            "governing_reference": "APP-SUP-OC-2022 revision 2.1; MAT-REF-2022-W52; weekly_order_board.xlsx",
        }
        resolved.append({field: row.get(field, "") for field in fields})
    return resolved


def table_rows_from_sheet(path: Path, sheet_name: str) -> tuple[list[str], list[dict[str, object]]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet = workbook[sheet_name]
    values = list(sheet.iter_rows(values_only=True))
    fields = [str(value) for value in values[0]]
    return fields, [dict(zip(fields, row)) for row in values[1:]]


def numeric_inputs() -> dict[str, tuple[list[str], list[dict[str, object]]]]:
    inputs: dict[str, tuple[list[str], list[dict[str, object]]]] = {}

    ppi_values: dict[int, list[float]] = defaultdict(list)
    with (DATA / "data" / "market" / "price_indicator_quarterly.csv").open(newline="") as handle:
        for row in csv.DictReader(handle):
            year = int(row["year"])
            if 2015 <= year <= 2024:
                ppi_values[year].append(float(row["ppi_value"]))
    ppi_fields = ["year", "annual_ppi_average", "observation_count", "source_release"]
    ppi_rows = [
        {
            "year": year,
            "annual_ppi_average": sum(values) / len(values),
            "observation_count": len(values),
            "source_release": "price_indicator_quarterly.csv",
        }
        for year, values in sorted(ppi_values.items())
    ]
    inputs["PPI_Annual"] = (ppi_fields, ppi_rows)

    macro_path = DATA / "data" / "macro" / "annual_context_panel.csv"
    with macro_path.open(newline="") as handle:
        macro_rows = list(csv.DictReader(handle))
    macro_fields = list(macro_rows[0])
    inputs["Macro_Context"] = (macro_fields, macro_rows)

    inputs["Procurement_Orders"] = table_rows_from_sheet(
        DATA / "data" / "procurement" / "weekly_order_board.xlsx", "Orders"
    )

    receipts_path = DATA / "data" / "procurement" / "receipt_extract.csv"
    with receipts_path.open(newline="") as handle:
        receipt_rows = list(csv.DictReader(handle))
    receipt_fields = list(receipt_rows[0])
    matcher_log = (DATA / "logs" / "fulfillment" / "2022-12-31.log").read_text()
    match = re.search(
        r"matched ASN (?P<receipt_id>\S+) order_ref=(?P<po>\S+) erp_item=(?P<erp>\S+) received_qty_kg=(?P<qty>-?\d+) status=(?P<status>\S+)",
        matcher_log,
    )
    if not match:
        raise ValueError("Late matcher receipt was not found")
    receipt_rows.append(
        {
            "receipt_id": match["receipt_id"],
            "po_number": match["po"],
            "erp_code": match["erp"],
            "receipt_date": "31/12/2022",
            "qty_kg": match["qty"],
            "status": match["status"],
            "source_note": "late WMS matcher order_ref resolved to po_number",
        }
    )
    inputs["Procurement_Receipts"] = (receipt_fields, receipt_rows)

    material_path = DATA / "config" / "procurement" / "material_catalog.json"
    material_rows = json.loads(material_path.read_text())["records"]
    material_fields = list(material_rows[0])
    inputs["Material_Catalog"] = (material_fields, material_rows)

    inputs["Supply_Pathways"] = table_rows_from_sheet(
        DATA / "data" / "supply" / "supply_pathways.xlsx", "Tier_Relationships"
    )

    rule_fields = ["rule_id", "parameter", "numeric_value", "unit", "effective_date", "governing_reference"]
    rule_rows = [
        {"rule_id": "RULE-THRESHOLD", "parameter": "review_threshold", "numeric_value": 35000, "unit": "GBP", "effective_date": "2023-01-27", "governing_reference": "REV-MR-23-03 section 3"},
        {"rule_id": "RULE-RESERVE", "parameter": "late_expedite_reserve", "numeric_value": 0.12, "unit": "ratio", "effective_date": "2022-12-20", "governing_reference": "FIN-RES-22-118"},
        {"rule_id": "RULE-NWD-1", "parameter": "non_working_date", "numeric_value": 20221226, "unit": "YYYYMMDD", "effective_date": "2022-09-29", "governing_reference": "OPS-CAL-2022"},
        {"rule_id": "RULE-NWD-2", "parameter": "non_working_date", "numeric_value": 20221227, "unit": "YYYYMMDD", "effective_date": "2022-09-29", "governing_reference": "OPS-CAL-2022"},
        {"rule_id": "RULE-NWD-3", "parameter": "non_working_date", "numeric_value": 20221228, "unit": "YYYYMMDD", "effective_date": "2022-12-22", "governing_reference": "OPS-CAL-2022 addendum"},
        {"rule_id": "RULE-NWD-4", "parameter": "non_working_date", "numeric_value": 20230102, "unit": "YYYYMMDD", "effective_date": "2022-09-29", "governing_reference": "OPS-CAL-2022"},
        {"rule_id": "RULE-WAIVER", "parameter": "transition_waiver_end", "numeric_value": 20230106, "unit": "YYYYMMDD", "effective_date": "2022-12-21", "governing_reference": "CR-2219"},
    ]
    inputs["Operational_Rules"] = (rule_fields, rule_rows)
    return inputs


def add_sheet(workbook: Workbook, title: str, fields: list[str], rows: list[dict[str, object]]) -> None:
    sheet = workbook.create_sheet(title)
    sheet.append(fields)
    for row in rows:
        sheet.append([row[field] for field in fields])
    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    table = Table(displayName=title.replace("_", ""), ref=sheet.dimensions)
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2", showFirstColumn=False, showLastColumn=False,
        showRowStripes=True, showColumnStripes=False
    )
    sheet.add_table(table)
    for column in sheet.columns:
        letter = column[0].column_letter
        width = min(max(len(str(cell.value or "")) for cell in column) + 2, 55)
        sheet.column_dimensions[letter].width = width


def build() -> None:
    OUTPUT.mkdir(exist_ok=True)
    schema = json.loads(SCHEMA_PATH.read_text())
    record_fields = schema["record_population_fields"]
    quality_fields = schema["quality_decision_fields"]
    governance_fields = schema["governance_fields"]
    entity_fields = schema["entity_resolution_fields"]
    records = catalog_rows(record_fields) + bridge_rows(record_fields)
    decisions = quality_decisions(quality_fields)
    governance = governance_rows(governance_fields)
    entities = entity_resolution_rows(entity_fields)
    analysis_inputs = numeric_inputs()

    workbook = Workbook()
    summary = workbook.active
    summary.title = "Read_Me"
    summary.append(["Weekly planning review register"])
    summary.append(["Purpose", "Pre-analysis source enumeration with explicit provenance and quality decisions"])
    summary.append(["Official catalog records", sum(r["source_classification"] == "official_catalog" for r in records)])
    summary.append(["Governed bridge records", sum(r["source_classification"] == "modeled_export_bridge" for r in records)])
    summary.append(["Quality decisions", len(decisions)])
    summary.append(["Governed parameters", len(governance)])
    summary.append(["Resolved canonical entities", len(entities)])
    summary.append(["Numeric input sheets", len(analysis_inputs)])
    summary.append(["2019 distinction", "Official catalog bundle is import-only; export is modeled"])
    summary.column_dimensions["A"].width = 28
    summary.column_dimensions["B"].width = 85
    summary["A1"].font = Font(size=14, bold=True)
    add_sheet(workbook, "Source_Records", record_fields, records)
    add_sheet(workbook, "Quality_Decisions", quality_fields, decisions)
    add_sheet(workbook, "Governance", governance_fields, governance)
    add_sheet(workbook, "Entity_Resolution", entity_fields, entities)
    for title, (fields, rows) in analysis_inputs.items():
        add_sheet(workbook, title, fields, rows)
    workbook.save(OUTPUT / "weekly_planning_review_register.xlsx")


if __name__ == "__main__":
    build()
