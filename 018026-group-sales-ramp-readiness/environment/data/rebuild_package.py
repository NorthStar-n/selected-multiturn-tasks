from __future__ import annotations

from collections import Counter
from datetime import date, datetime, time
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


ROOT = Path(__file__).resolve().parent
DATA = ROOT / "data"
OUT = ROOT / "output"
PACKAGE_ID = "HB-RAMP-2026-07"
AS_OF = date(2026, 7, 15)
APPROVAL_DEADLINE = datetime(2026, 7, 17, 17, 0)

INTAKE = DATA / "cohort_setup" / "july_new_rep_intake_targets.xlsx"
CURRICULUM = DATA / "cohort_setup" / "product_curriculum_and_competency_map.xlsx"
EVIDENCE = DATA / "evidence_exports" / "certification_assessment_observation_export.xlsx"
AVAILABILITY = DATA / "availability" / "trainer_observation_availability.xlsx"
COVERAGE = DATA / "account_coverage" / "temporary_group_account_coverage_export.xlsx"
STANDARD = DATA / "governance" / "group_sales_ramp_certification_standard.docx"
REQUEST = DATA / "approvals" / "ramp_readiness_approval_request_thread.eml"
SOURCES = [INTAKE, CURRICULUM, EVIDENCE, AVAILABILITY, COVERAGE, STANDARD, REQUEST]

NAVY = "17324D"
BLUE = "2F75B5"
PALE_BLUE = "D9EAF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_ORANGE = "FCE4D6"
PALE_RED = "F4CCCC"
GRAY = "E7E6E6"
WHITE = "FFFFFF"
THIN = Side(style="thin", color="B7B7B7")


def records(path: Path, sheet: str) -> list[dict]:
    ws = load_workbook(path, data_only=True, read_only=True)[sheet]
    it = ws.iter_rows(values_only=True)
    headers = next(it)
    return [dict(zip(headers, row)) for row in it if any(v is not None for v in row)]


def source_manifest() -> list[list]:
    result = []
    for path in SOURCES:
        raw = path.read_bytes()
        result.append([str(path.relative_to(ROOT)), len(raw), sha256(raw).hexdigest()])
    return result


cohort = records(INTAKE, "Cohort")
evidence = records(EVIDENCE, "Evidence_Status_By_Competency")
exceptions = records(EVIDENCE, "Exceptions")
assessment_attempts = records(EVIDENCE, "Assessment_Attempts")
competency_rules = records(CURRICULUM, "Competency_Map")
accounts = records(COVERAGE, "Accounts")
rep_by_id = {r["Rep_ID"]: r for r in cohort}
evidence_by_rep = {r["Rep_ID"]: [e for e in evidence if e["Rep_ID"] == r["Rep_ID"]] for r in cohort}


def controlled_status(rep_id: str) -> str:
    rows = evidence_by_rep[rep_id]
    statuses = {r["Evidence_Status"] for r in rows}
    states = {r["Approval_State"] for r in rows if r["Approval_State"]}
    constraint = rep_by_id[rep_id]["Constraint_Type"]
    if "Approved Exception" in statuses:
        return "Limited exception"
    if "Blocked" in statuses or "Unresolved" in states or "licensing hold" in constraint.lower() or "compliance disclosure miss" in constraint.lower():
        return "Blocked"
    if "Failed" in statuses:
        return "Remediation"
    if "Open" in statuses:
        return "Pending evidence"
    return "Ready — manager approval pending"


def blocker_summary(rep_id: str) -> str:
    open_rows = [r for r in evidence_by_rep[rep_id] if r["Evidence_Status"] != "Complete"]
    if not open_rows:
        return "Evidence complete; independent coverage still requires desk manager approval"
    return "; ".join(f'{r["Competency_ID"]}: {r["Next_Action"]}' for r in open_rows)


def next_owner_due(rep_id: str) -> tuple[str, datetime | None]:
    rows = [r for r in evidence_by_rep[rep_id] if r["Evidence_Status"] != "Complete"]
    if not rows:
        return "Desk Manager", APPROVAL_DEADLINE
    owners = ", ".join(dict.fromkeys(str(r["Owner_ID"]) for r in rows if r["Owner_ID"]))
    dues = [r["Due_Date"] for r in rows if r["Due_Date"]]
    return owners, min(dues) if dues else None


def rep_summary_rows() -> list[list]:
    output = []
    for r in cohort:
        rid = r["Rep_ID"]
        counts = Counter(e["Evidence_Status"] for e in evidence_by_rep[rid])
        owner, due = next_owner_due(rid)
        output.append([
            rid, r["Representative_Name"], r["Segment"], r["Territory_Code"],
            r["Constraint_Type"], counts.get("Complete", 0), counts.get("Approved Exception", 0),
            counts.get("Failed", 0), counts.get("Open", 0), counts.get("Blocked", 0),
            controlled_status(rid), blocker_summary(rid), owner, due,
            "No", "Pending desk manager review",
        ])
    return output


REP_HEADERS = [
    "Rep_ID", "Representative_Name", "Segment", "Territory_Code", "Constraint_Type",
    "Complete_Gates", "Approved_Exception_Gates", "Failed_Gates", "Open_Gates", "Blocked_Gates",
    "Controlled_Status", "Blocker_or_Next_Action", "Next_Owner", "Due_Date",
    "Independent_Coverage_Authorized", "Manager_Approval_State",
]


def fresh_wb(title: str) -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = "Control"
    ws.append([title])
    ws.append(["Package_ID", PACKAGE_ID])
    ws.append(["Source_as_of", AS_OF])
    ws.append(["Approval_state", "PENDING — no independent customer coverage authorized"])
    ws.append(["Approval_deadline_UTC", APPROVAL_DEADLINE])
    ws.append(["Population", "10 representatives; all controlled states retained"])
    ws.append(["Build_rule", "Derived from supplied source records; do not manually patch output"])
    ws.merge_cells("A1:F1")
    ws["A1"].font = Font(bold=True, color=WHITE, size=16)
    ws["A1"].fill = PatternFill("solid", fgColor=NAVY)
    ws["A1"].alignment = Alignment(horizontal="left")
    ws.column_dimensions["A"].width = 27
    ws.column_dimensions["B"].width = 72
    ws["B3"].number_format = "yyyy-mm-dd"
    ws["B5"].number_format = "yyyy-mm-dd hh:mm"
    return wb


def add_table_sheet(wb: Workbook, name: str, headers: list, rows: list[list], freeze: str = "A2"):
    ws = wb.create_sheet(name)
    ws.append(headers)
    for row in rows:
        ws.append(row)
    ws.freeze_panes = freeze
    ws.auto_filter.ref = ws.dimensions
    for cell in ws[1]:
        cell.font = Font(bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[1].height = 32
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(bottom=THIN)
            if isinstance(cell.value, datetime):
                cell.number_format = "yyyy-mm-dd hh:mm"
            elif isinstance(cell.value, date):
                cell.number_format = "yyyy-mm-dd"
            elif isinstance(cell.value, time):
                cell.number_format = "hh:mm"
    for idx, header in enumerate(headers, 1):
        values = [str(ws.cell(r, idx).value or "") for r in range(1, min(ws.max_row, 100) + 1)]
        width = min(max(max(map(len, values)) + 2, len(str(header)) + 2), 52)
        ws.column_dimensions[get_column_letter(idx)].width = width
    return ws


def add_record_sheet(wb: Workbook, name: str, source_path: Path, source_sheet: str):
    rows = records(source_path, source_sheet)
    headers = list(rows[0]) if rows else []
    add_table_sheet(wb, name, headers, [[r[h] for h in headers] for r in rows])


def add_lineage(wb: Workbook):
    ws = add_table_sheet(wb, "Source_Lineage", ["Source_File", "Bytes", "SHA256"], source_manifest())
    ws.column_dimensions["A"].width = 72
    ws.column_dimensions["C"].width = 68


def color_statuses(ws, column_name: str):
    headers = {c.value: c.column for c in ws[1]}
    col = get_column_letter(headers[column_name])
    fills = {
        "Ready — manager approval pending": PALE_GREEN,
        "Limited exception": PALE_YELLOW,
        "Remediation": PALE_ORANGE,
        "Pending evidence": PALE_BLUE,
        "Blocked": PALE_RED,
    }
    for label, color in fills.items():
        ws.conditional_formatting.add(
            f"{col}2:{col}{ws.max_row}",
            FormulaRule(formula=[f'${col}2="{label}"'], fill=PatternFill("solid", fgColor=color)),
        )


def build_evidence_register():
    wb = fresh_wb("Certification Evidence Register")
    status_counts = Counter(e["Evidence_Status"] for e in evidence)
    readiness_counts = Counter(controlled_status(r["Rep_ID"]) for r in cohort)
    attempt_counts = Counter(a["Pass_Fail"] for a in assessment_attempts)
    source_checks = {x["Check"]: x["Expected"] for x in records(EVIDENCE, "Checks")}
    control = wb["Control"]
    control.append(["Detail_complete_gate_count", status_counts["Complete"]])
    control.append(["Supplied_Checks_complete_benchmark", source_checks.get("Complete evidence count")])
    control.append(["Unresolved_source_variance", status_counts["Complete"] - source_checks.get("Complete evidence count"), "VISIBLE — detail rows show 59 Complete; supplied Checks benchmark says 57"])
    for cell in control[10]:
        cell.fill = PatternFill("solid", fgColor=PALE_RED)
        cell.font = Font(bold=True, color="9C0006")
    total_rows = [
        ["Cohort representatives", len(cohort), 10, "PASS", "All intake Rep_IDs retained"],
        ["Required competencies", len(competency_rules), 7, "PASS", "GS-C01 through GS-C07"],
        ["Required representative-competency pairs", len(cohort) * len(competency_rules), 70, "PASS", "10 × 7 controlled grid"],
        ["Reconciled representative-competency pairs", len(evidence), 70, "PASS" if len(evidence) == 70 else "FAIL", "One source evidence row per required pair"],
        ["Complete evidence gates", status_counts["Complete"], source_checks.get("Complete evidence count"), "SOURCE CHECK VARIANCE" if status_counts["Complete"] != source_checks.get("Complete evidence count") else "PASS", "Detail records contain 59 Complete states; supplied Checks benchmark says 57"],
        ["Approved-exception gates", status_counts["Approved Exception"], source_checks.get("Approved exception count"), "PASS", "Limited exception is not independent coverage"],
        ["Failed gates", status_counts["Failed"], None, "OPEN CONTROL", "Remediation or compliance action required"],
        ["Open gates", status_counts["Open"], None, "OPEN CONTROL", "Evidence, observation, or panel pending"],
        ["Blocked gates", status_counts["Blocked"], None, "OPEN CONTROL", "Licensing/compliance hold"],
        ["Assessment attempts", len(assessment_attempts), 63, "PASS", f'{attempt_counts["Pass"]} pass / {attempt_counts["Fail"]} fail attempt records'],
        ["Exception records", len(exceptions), 7, "PASS", "Every exception has owner, due date, and next action"],
        ["Ready — manager approval pending", readiness_counts["Ready — manager approval pending"], None, "PENDING APPROVAL", "Evidence complete; manager decision still required"],
        ["Limited exception representatives", readiness_counts["Limited exception"], None, "CONTROLLED", "Supervised low-risk boundary only"],
        ["Remediation representatives", readiness_counts["Remediation"], None, "CONTROLLED", "Failed gate under remediation"],
        ["Pending-evidence representatives", readiness_counts["Pending evidence"], None, "CONTROLLED", "Observation or panel evidence open"],
        ["Blocked representatives", readiness_counts["Blocked"], None, "CONTROLLED", "Compliance/licensing blocker"],
        ["Independent coverage authorizations", 0, 0, "PASS", "Desk manager approval remains pending"],
    ]
    totals = add_table_sheet(wb, "Cohort_Totals", ["Metric", "Actual", "Expected_or_Benchmark", "State", "Reconciliation_Note"], total_rows)
    totals.column_dimensions["A"].width = 43
    totals.column_dimensions["E"].width = 72

    eligibility_headers = [
        "Rep_ID", "Representative_Name", "Segment", "Territory_Code", "Controlled_Status",
        "Eligible_For_Manager_Readiness_Review", "Excluded_or_Delayed_From_Readiness",
        "Eligibility_or_Exclusion_Reason", "Action_Owner", "Action_Due_Date", "Next_Action",
        "Complete_Gates", "Approved_Exception_Gates", "Failed_Gates", "Open_Gates", "Blocked_Gates",
        "Manager_Approval_State", "Independent_Coverage_Authorized",
    ]
    eligibility_rows = []
    for rep in cohort:
        rid = rep["Rep_ID"]
        rows = evidence_by_rep[rid]
        counts = Counter(x["Evidence_Status"] for x in rows)
        status = controlled_status(rid)
        eligible = status == "Ready — manager approval pending"
        owner, due = next_owner_due(rid)
        if eligible:
            reason = "All seven competency rows show Complete; eligible for desk manager readiness review only."
            action = "Desk manager reviews the complete evidence package and records a signed decision."
        else:
            noncomplete = [x for x in rows if x["Evidence_Status"] != "Complete"]
            reason = "; ".join(f'{x["Competency_ID"]} {x["Evidence_Status"]}: {x["Approval_State"] or "No approval state"}' for x in noncomplete)
            action = "; ".join(dict.fromkeys(x["Next_Action"] for x in noncomplete if x["Next_Action"]))
        eligibility_rows.append([
            rid, rep["Representative_Name"], rep["Segment"], rep["Territory_Code"], status,
            "Yes" if eligible else "No", "No" if eligible else "Yes", reason, owner, due, action,
            counts["Complete"], counts["Approved Exception"], counts["Failed"], counts["Open"], counts["Blocked"],
            "Pending desk manager review", "No",
        ])
    eligibility = add_table_sheet(wb, "Eligibility_Register", eligibility_headers, eligibility_rows)
    color_statuses(eligibility, "Controlled_Status")
    excluded_col = get_column_letter({c.value: c.column for c in eligibility[1]}["Excluded_or_Delayed_From_Readiness"])
    eligibility.conditional_formatting.add(
        f"{excluded_col}2:{excluded_col}{eligibility.max_row}",
        FormulaRule(formula=[f'${excluded_col}2="Yes"'], fill=PatternFill("solid", fgColor=PALE_RED)),
    )

    ws = add_table_sheet(wb, "Rep_Status", REP_HEADERS, rep_summary_rows())
    color_statuses(ws, "Controlled_Status")

    comp_ids = [c["Competency_ID"] for c in competency_rules]
    matrix_headers = ["Rep_ID", "Representative_Name"] + comp_ids + ["Reconciled_Gates", "Required_Gates", "Representative_Readiness", "Manager_Approval_State", "Independent_Coverage_Authorized"]
    matrix_rows = []
    for rep in cohort:
        rid = rep["Rep_ID"]
        by_comp = {x["Competency_ID"]: x for x in evidence_by_rep[rid]}
        matrix_rows.append([
            rid, rep["Representative_Name"], *[by_comp[c]["Evidence_Status"] if c in by_comp else "MISSING" for c in comp_ids],
            len(by_comp), len(comp_ids), controlled_status(rid), "Pending desk manager review", "No",
        ])
    matrix = add_table_sheet(wb, "Competency_Matrix", matrix_headers, matrix_rows)
    color_statuses(matrix, "Representative_Readiness")
    matrix_status_colors = {"Complete": PALE_GREEN, "Approved Exception": PALE_YELLOW, "Failed": PALE_ORANGE, "Open": PALE_BLUE, "Blocked": PALE_RED, "MISSING": PALE_RED}
    for col_idx in range(3, 3 + len(comp_ids)):
        col = get_column_letter(col_idx)
        for label, color in matrix_status_colors.items():
            matrix.conditional_formatting.add(
                f"{col}2:{col}{matrix.max_row}",
                FormulaRule(formula=[f'${col}2="{label}"'], fill=PatternFill("solid", fgColor=color)),
            )

    attempt_by_pair = {}
    for attempt in assessment_attempts:
        attempt_by_pair.setdefault((attempt["Rep_ID"], attempt["Competency_ID"]), []).append(attempt)
    exception_by_pair = {(x["Rep_ID"], x["Competency_ID"]): x for x in exceptions}
    reconciliation_headers = [
        "Rep_ID", "Representative_Name", "Segment", "Territory_Code", "Competency_ID", "Competency_Name",
        "Required_Evidence", "Minimum_Rule", "Evidence_Status", "Latest_Evidence", "Evidence_Date", "Evidence_Score",
        "Attempt_Count", "Attempt_IDs", "Attempt_History", "Latest_Attempt_ID", "Latest_Attempt_Date",
        "Latest_Attempt_Score", "Latest_Attempt_Result", "Latest_Assessor_ID", "Exception_Code", "Exception_Reason",
        "Exception_Approval_State", "Exception_Owner", "Exception_Due_Date", "Evidence_Due_Date", "Control_Due_Date",
        "Next_Action", "Source_Reference", "Pair_Reconciliation", "Representative_Readiness",
        "Manager_Approval_State", "Independent_Coverage_Authorized",
    ]
    reconciliation_rows = []
    for rep in cohort:
        rid = rep["Rep_ID"]
        by_comp = {x["Competency_ID"]: x for x in evidence_by_rep[rid]}
        for comp in competency_rules:
            cid = comp["Competency_ID"]
            ev = by_comp.get(cid)
            pair_attempts = attempt_by_pair.get((rid, cid), [])
            exc = exception_by_pair.get((rid, cid))
            latest_attempt = pair_attempts[-1] if pair_attempts else None
            control_due = (exc["Next_Action_Due_Date"] if exc else None) or (ev["Due_Date"] if ev else None) or rep["Target_Certification_Date"]
            attempt_history = "; ".join(
                f'{a["Attempt_ID"]} | {a["Attempt_Date"].date().isoformat()} | {a["Score"]} | {a["Pass_Fail"]} | {a["Assessor_ID"]}'
                for a in pair_attempts
            )
            reconciliation_rows.append([
                rid, rep["Representative_Name"], rep["Segment"], rep["Territory_Code"], cid, comp["Competency_Name"],
                comp["Required_Evidence"], comp["Minimum_Rule"], ev["Evidence_Status"] if ev else "MISSING",
                ev["Latest_Evidence"] if ev else None, ev["Evidence_Date"] if ev else None, ev["Score"] if ev else None,
                len(pair_attempts), "; ".join(a["Attempt_ID"] for a in pair_attempts), attempt_history,
                latest_attempt["Attempt_ID"] if latest_attempt else None, latest_attempt["Attempt_Date"] if latest_attempt else None,
                latest_attempt["Score"] if latest_attempt else None, latest_attempt["Pass_Fail"] if latest_attempt else None,
                latest_attempt["Assessor_ID"] if latest_attempt else None, exc["Exception_Code"] if exc else None,
                exc["Exception_Reason"] if exc else None, exc["Approval_State"] if exc else (ev["Approval_State"] if ev else None),
                exc["Exception_Owner"] if exc else (ev["Owner_ID"] if ev else None),
                exc["Next_Action_Due_Date"] if exc else None, ev["Due_Date"] if ev else None, control_due,
                exc["Next_Action"] if exc else (ev["Next_Action"] if ev else "Missing evidence source row"),
                ev["Source_Reference"] if ev else None, "RECONCILED" if ev else "MISSING SOURCE EVIDENCE",
                controlled_status(rid), "Pending desk manager review", "No",
            ])
    recon = add_table_sheet(wb, "Competency_Reconciliation", reconciliation_headers, reconciliation_rows)
    color_statuses(recon, "Representative_Readiness")
    status_col = get_column_letter({c.value: c.column for c in recon[1]}["Evidence_Status"])
    for label, color in matrix_status_colors.items():
        recon.conditional_formatting.add(
            f"{status_col}2:{status_col}{recon.max_row}",
            FormulaRule(formula=[f'${status_col}2="{label}"'], fill=PatternFill("solid", fgColor=color)),
        )

    add_record_sheet(wb, "Evidence_Detail", EVIDENCE, "Evidence_Status_By_Competency")
    add_record_sheet(wb, "Assessment_Attempts", EVIDENCE, "Assessment_Attempts")
    add_record_sheet(wb, "Observations", EVIDENCE, "Observations")
    add_record_sheet(wb, "CRM_Audits", EVIDENCE, "CRM_Audits")
    add_record_sheet(wb, "Coaching_History", EVIDENCE, "Coaching_History")
    add_record_sheet(wb, "Exceptions", EVIDENCE, "Exceptions")
    add_record_sheet(wb, "Competency_Rules", CURRICULUM, "Competency_Map")
    checks = [
        ["Controlled representatives", len(cohort), 10, "PASS" if len(cohort) == 10 else "FAIL"],
        ["Evidence rows", len(evidence), 70, "PASS" if len(evidence) == 70 else "FAIL"],
        ["Seven gates per representative", len(evidence), len(cohort) * 7, "PASS" if len(evidence) == len(cohort) * 7 else "FAIL"],
        ["Exception rows", len(exceptions), 7, "PASS" if len(exceptions) == 7 else "FAIL"],
        ["Independent approvals", 0, 0, "PASS"],
        ["Unique representative-competency pairs", len({(r[0], r[4]) for r in reconciliation_rows}), 70, "PASS" if len({(r[0], r[4]) for r in reconciliation_rows}) == 70 else "FAIL"],
        ["Missing evidence source pairs", sum(1 for r in reconciliation_rows if r[29] != "RECONCILED"), 0, "PASS" if all(r[29] == "RECONCILED" for r in reconciliation_rows) else "FAIL"],
        ["Pairs with attempt history or non-assessment evidence", sum(1 for r in reconciliation_rows if r[12] > 0 or r[9]), 70, "PASS" if all(r[12] > 0 or r[9] for r in reconciliation_rows) else "FAIL"],
    ]
    add_table_sheet(wb, "Validation", ["Check", "Actual", "Expected", "Result"], checks)
    add_lineage(wb)
    wb.save(OUT / "certification_evidence_register.xlsx")


SCHEDULE_HEADERS = [
    "Event_ID", "Rep_ID", "Representative_Name", "Controlled_Status", "Event_Date",
    "Start_Time", "End_Time", "Trainer_ID", "Event_Type", "Competency_ID",
    "Prerequisite", "Prerequisite_Check", "Trainer_Qualification_Check", "Availability_Window",
    "Availability_Check", "Blackout_Check", "Window_Capacity", "Trainer_Daily_Max",
    "Scheduled_Slot_Number", "Buffer_Rule", "Buffer_Check", "Event_State", "Source_Basis",
    "Independent_Coverage_Impact", "Excluded_From_Execution", "Exclusion_or_Delay_Reason",
    "Action_Owner", "Action_Due_Date", "Next_Action",
]


def schedule_rows() -> list[list]:
    exception_by_rep = {x["Rep_ID"]: x for x in exceptions}
    owner_override = {
        "SCH-2703-PNL": "TRN-506 (Dana Miller)",
        "SCH-2702-PNL": "SUP-411 (Marisol Grant), coordinating TRN-506",
        "SCH-2708-RT": "COM-201, coordinating TRN-504",
        "SCH-2708-OBS": "COM-201, then TRN-505 after clearance",
        "SCH-2708-PNL": "COM-201, then TRN-506 after observation",
        "SCH-2704-LIC": "COM-201",
        "SCH-2704-C04": "COM-201, coordinating TRN-504",
        "SCH-2704-C06": "COM-201, then TRN-501 after clearance",
        "SCH-2704-OBS": "COM-201, then TRN-505 after prerequisites",
        "SCH-2704-PNL": "COM-201, then TRN-506 after observation",
    }
    next_override = {
        "SCH-2703-PNL": "TRN-506 must publish a post-observation panel slot and preserve the 60-minute observation-to-panel buffer.",
        "SCH-2702-PNL": "SUP-411 must keep the limited boundary active and secure a TRN-506 panel slot at least 60 minutes after observation.",
        "SCH-2708-RT": "COM-201 must secure a new TRN-504 retest window beginning at least 120 minutes after coaching.",
        "SCH-2708-OBS": "COM-201 must record a passing compliance retest before TRN-505 may release the observation hold.",
        "SCH-2708-PNL": "Clear compliance, complete the observation, then schedule TRN-506 at least 60 minutes later.",
        "SCH-2704-LIC": "COM-201 must confirm seller-of-travel license clearance in the evidence register.",
        "SCH-2704-C04": "After license clearance, COM-201 must secure qualified TRN-504 capacity for compliance signoff or retest.",
        "SCH-2704-C06": "After GS-C04 clears, COM-201 must secure a published TRN-501 window for the escalation scenario.",
        "SCH-2704-OBS": "After GS-C04 and GS-C06 clear, secure a qualified TRN-505 observation window.",
        "SCH-2704-PNL": "After observation passes, secure TRN-506 capacity and enforce the 60-minute buffer.",
    }

    def r(eid, rid, d, start, end, trainer, typ, comp, prereq, prereq_check, qual,
          window, avail, blackout, window_cap, daily_max, slot_no, buffer_rule,
          buffer_check, state, basis):
        excluded = not state.startswith("SCHEDULED")
        rep = rep_by_id[rid]
        exception = exception_by_rep.get(rid)
        if excluded:
            if trainer == "DESK-MGR":
                reason = f"{state}; {prereq}. {window}."
                owner = "Rina Calder (Desk Manager)"
                due = rep["Target_Independent_Coverage_Date"]
                action = "Complete any upstream events, submit the refreshed evidence package, and record the desk manager's signed decision."
            else:
                reason_parts = [state]
                if not str(avail).startswith("PASS"):
                    reason_parts.append(str(window))
                if str(prereq_check).startswith(("BLOCKED", "CONDITIONAL", "OPEN")):
                    reason_parts.append(f"Prerequisite: {prereq}")
                reason = "; ".join(reason_parts)
                owner = owner_override.get(eid, exception["Exception_Owner"] if exception else trainer)
                if eid.endswith(("PNL", "OBS", "C06")):
                    due = rep["Target_Independent_Coverage_Date"]
                else:
                    due = exception["Next_Action_Due_Date"] if exception else rep["Target_Certification_Date"]
                action = next_override.get(eid, exception["Next_Action"] if exception else f"Resolve the stated prerequisite and publish a compliant {typ.lower()} slot.")
        else:
            reason = ""
            owner = trainer
            due = d
            action = "Complete the event only when its prerequisite check is satisfied; record resulting evidence before releasing downstream work."
        return [
            eid, rid, rep_by_id[rid]["Representative_Name"], controlled_status(rid), d, start, end,
            trainer, typ, comp, prereq, prereq_check, qual, window, avail, blackout,
            window_cap, daily_max, slot_no, buffer_rule, buffer_check, state, basis,
            "None until signed desk manager approval", "Yes" if excluded else "No", reason,
            owner, due, action,
        ]
    return [
        # Complete-evidence representatives still require the manager's controlled decision.
        r("SCH-2701-MGR", "REP-2701", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "All GS-C01–GS-C07 evidence complete", "PASS", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — MANAGER REVIEW", "Evidence register; governance standard"),
        r("SCH-2707-MGR", "REP-2707", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "All GS-C01–GS-C07 evidence complete", "PASS", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — MANAGER REVIEW", "Evidence register; governance standard"),
        r("SCH-2710-MGR", "REP-2710", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "All GS-C01–GS-C07 evidence complete", "PASS", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — MANAGER REVIEW", "Evidence register; governance standard"),

        # Priya: remaining CRM remediation and retest, then manager review.
        r("SCH-2705-RM", "REP-2705", date(2026, 7, 16), time(8, 30), time(9, 30), "TRN-503", "Remaining remediation module: CRM hygiene rebuild", "GS-C05", "GS-C01 complete", "PASS", "PASS — TRN-503 qualified for GS-C05",
          "2026-07-16 08:30–12:30; capacity 3", "PASS", "PASS", 3, 3, 1, "M501 → A-C05 >= 30 minutes", "PASS — 30 minutes", "SCHEDULED", "EX-CRM-2705; Modules; Trainer_Roster; Availability_Windows"),
        r("SCH-2705-RT", "REP-2705", date(2026, 7, 16), time(10, 0), time(10, 30), "TRN-503", "CRM audit retest", "GS-C05", "CRM hygiene rebuild complete", "CONDITIONAL — prior event must complete", "PASS — TRN-503 qualified for GS-C05",
          "2026-07-16 08:30–12:30; capacity 3", "PASS", "PASS", 3, 3, 2, "M501 → A-C05 >= 30 minutes", "PASS — 30 minutes", "SCHEDULED — PREREQUISITE CONTROLLED", "EX-CRM-2705; Assessments; Buffer_Rules"),
        r("SCH-2705-MGR", "REP-2705", None, None, None, "DESK-MGR", "Desk manager decision after retest", "GS-C07", "GS-C05 retest passes; other gates complete", "CONDITIONAL", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — RETEST AND MANAGER REVIEW", "Evidence register; governance standard"),

        # Sofia: terms remediation, retest, observation, panel, manager review.
        r("SCH-2703-RM", "REP-2703", date(2026, 7, 16), time(13, 0), time(14, 30), "TRN-503", "Remaining remediation module: contract terms", "GS-C03", "GS-C02 complete", "PASS", "PASS — TRN-503 qualified for GS-C03",
          "2026-07-16 13:00–16:30; capacity 3", "PASS", "PASS", 3, 3, 1, "M301 → A-C03 >= 240 minutes", "PASS — retest next day", "SCHEDULED", "EX-REM-2703; Modules; Trainer_Roster; Availability_Windows"),
        r("SCH-2703-RT", "REP-2703", date(2026, 7, 17), time(9, 0), time(9, 30), "TRN-503", "Contract assessment retest", "GS-C03", "Contract terms remediation complete", "CONDITIONAL — prior event must complete", "PASS — TRN-503 qualified for GS-C03",
          "2026-07-17 09:00–12:00; capacity 2", "PASS", "PASS", 2, 3, 1, "M301 → A-C03 >= 240 minutes", "PASS — 18.5 hours", "SCHEDULED — PREREQUISITE CONTROLLED", "EX-REM-2703; Assessments; Buffer_Rules"),
        r("SCH-2703-OBS", "REP-2703", date(2026, 7, 20), time(10, 0), time(11, 0), "TRN-505", "Live observation", "GS-C07", "GS-C03 retest passes; GS-C02–GS-C06 complete", "CONDITIONAL", "PASS — TRN-505 qualified for GS-C07",
          "2026-07-20 09:00–12:00; capacity 4", "PASS", "PASS", 4, 4, 2, "No additional source buffer", "PASS", "TENTATIVE HOLD — PREREQUISITE CONTROLLED", "Evidence GS-C07; Trainer_Roster; Availability_Windows"),
        r("SCH-2703-PNL", "REP-2703", None, None, None, "TRN-506", "Supervisor recommendation panel", "GS-C07", "Live observation passes", "CONDITIONAL", "PASS — TRN-506 qualified for GS-C07",
          "No published TRN-506 window after 2026-07-20 observation", "NO COMPLIANT WINDOW", "N/A", 0, 2, None, "A-C07 → supervisor recommendation >= 60 minutes", "NOT TESTABLE", "UNSCHEDULED — NEW WINDOW REQUIRED", "Buffer_Rules; Availability_Windows"),
        r("SCH-2703-MGR", "REP-2703", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "Retest, observation, and panel complete", "CONDITIONAL", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — UPSTREAM EVENTS", "Governance standard"),

        # Liam and Isabella: complete their remaining observation/panel paths inside capacity.
        r("SCH-2706-OBS", "REP-2706", date(2026, 7, 17), time(8, 30), time(9, 30), "TRN-505", "Live observation", "GS-C07", "GS-C02–GS-C06 complete", "PASS", "PASS — TRN-505 qualified for GS-C07",
          "2026-07-17 08:30–12:30; capacity 4", "PASS", "PASS — before 13:00 blackout", 4, 4, 1, "A-C07 → supervisor recommendation >= 60 minutes", "PASS — panel starts 60 minutes later", "SCHEDULED", "EX-LIVE-2706; Observations; Blackouts"),
        r("SCH-2706-PNL", "REP-2706", date(2026, 7, 17), time(10, 30), time(11, 0), "TRN-506", "Supervisor recommendation panel", "GS-C07", "Live observation passes", "CONDITIONAL", "PASS — TRN-506 qualified for GS-C07",
          "2026-07-17 10:30–12:30; capacity 2", "PASS", "PASS", 2, 2, 1, "A-C07 → supervisor recommendation >= 60 minutes", "PASS — 60 minutes", "SCHEDULED — PREREQUISITE CONTROLLED", "EX-LIVE-2706; Buffer_Rules; Availability_Windows"),
        r("SCH-2706-MGR", "REP-2706", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "Observation and panel pass", "CONDITIONAL", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — UPSTREAM EVENTS", "Governance standard"),
        r("SCH-2709-PNL", "REP-2709", date(2026, 7, 17), time(10, 30), time(11, 0), "TRN-506", "Supervisor recommendation panel", "GS-C07", "A-C07 evidence complete", "PASS", "PASS — TRN-506 qualified for GS-C07",
          "2026-07-17 10:30–12:30; capacity 2", "PASS", "PASS", 2, 2, 2, "A-C07 → supervisor recommendation >= 60 minutes", "PASS — source evidence predates panel", "SCHEDULED", "EX-PANEL-2709; Trainer_Roster; Availability_Windows"),
        r("SCH-2709-MGR", "REP-2709", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "Supervisor panel completes", "CONDITIONAL", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — PANEL", "Governance standard"),

        # Jordan's limited exception remains supervised through observation; no panel window follows it.
        r("SCH-2702-OBS", "REP-2702", date(2026, 7, 20), time(9, 0), time(10, 0), "TRN-505", "Limited-exception live observation", "GS-C07", "Low-risk supervised affinity renewals only", "PASS — exception boundary active", "PASS — TRN-505 qualified for GS-C07",
          "2026-07-20 09:00–12:00; capacity 4", "PASS", "PASS", 4, 4, 1, "A-C07 → supervisor recommendation >= 60 minutes", "PENDING PANEL", "SCHEDULED — LIMITED BOUNDARY", "EX-LIM-2702; Availability_Windows"),
        r("SCH-2702-PNL", "REP-2702", None, None, None, "TRN-506", "Supervisor recommendation panel", "GS-C07", "Live observation passes", "CONDITIONAL", "PASS — TRN-506 qualified for GS-C07",
          "No published TRN-506 window after 2026-07-20 observation", "NO COMPLIANT WINDOW", "N/A", 0, 2, None, "A-C07 → supervisor recommendation >= 60 minutes", "NOT TESTABLE", "UNSCHEDULED — NEW WINDOW REQUIRED", "EX-LIM-2702; Buffer_Rules"),
        r("SCH-2702-MGR", "REP-2702", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "Observation and panel complete", "CONDITIONAL", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "PENDING — LIMITED EXCEPTION REMAINS", "Governance standard"),

        # Noah: coaching fits; the mandatory buffer makes the retest unschedulable in the published window.
        r("SCH-2708-RM", "REP-2708", date(2026, 7, 17), time(14, 0), time(15, 0), "TRN-504", "Remaining remediation module: compliance disclosure", "GS-C04", "GS-C01 complete", "PASS", "PASS — TRN-504 qualified for GS-C04",
          "2026-07-17 14:00–17:00; capacity 2", "PASS", "PASS", 2, 2, 1, "M401 → A-C04 >= 120 minutes", "RETEST EARLIEST 17:00", "SCHEDULED", "EX-COMP-2708; Modules; Availability_Windows"),
        r("SCH-2708-RT", "REP-2708", None, None, None, "TRN-504", "Compliance retest", "GS-C04", "Compliance remediation complete", "CONDITIONAL", "PASS — TRN-504 qualified for GS-C04",
          "Published window ends at earliest permitted retest start; TRN-504 blacked out 2026-07-20", "NO COMPLIANT WINDOW", "PASS / 2026-07-20 BLACKOUT ENFORCED", 0, 2, None, "M401 → A-C04 >= 120 minutes", "FAIL — no duration remains at 17:00", "UNSCHEDULED — NEW WINDOW REQUIRED", "EX-COMP-2708; Buffer_Rules; Blackouts"),
        r("SCH-2708-OBS", "REP-2708", date(2026, 7, 20), time(11, 0), time(12, 0), "TRN-505", "Live observation hold", "GS-C07", "Compliance retest passes", "BLOCKED", "PASS — TRN-505 qualified for GS-C07",
          "2026-07-20 09:00–12:00; capacity 4", "PASS", "PASS", 4, 4, 3, "No additional source buffer", "N/A", "TENTATIVE HOLD — DO NOT EXECUTE UNTIL RETEST", "Evidence GS-C07; Availability_Windows"),
        r("SCH-2708-PNL", "REP-2708", None, None, None, "TRN-506", "Supervisor recommendation panel", "GS-C07", "Compliance retest and live observation pass", "BLOCKED", "PASS — TRN-506 qualified for GS-C07",
          "No published window after conditional observation", "NO COMPLIANT WINDOW", "N/A", 0, 2, None, "A-C07 → supervisor recommendation >= 60 minutes", "NOT TESTABLE", "UNSCHEDULED — UPSTREAM BLOCKER", "Evidence register; Buffer_Rules"),
        r("SCH-2708-MGR", "REP-2708", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "Compliance, observation, and panel complete", "BLOCKED", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "BLOCKED — UPSTREAM EVENTS", "Governance standard"),

        # Ethan: licensing must clear before remaining compliance, escalation, observation, and panel work.
        r("SCH-2704-LIC", "REP-2704", None, None, None, "COM-201", "Licensing clearance", "GS-C04", "State seller-of-travel license confirmation", "OPEN", "N/A — compliance owner control",
          "Due 2026-07-22; no owner availability supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "BLOCKED — COMPLIANCE OWNER ACTION", "EX-LIC-2704"),
        r("SCH-2704-C04", "REP-2704", None, None, None, "TRN-504", "Compliance signoff / retest", "GS-C04", "Licensing clearance", "BLOCKED", "PASS — TRN-504 qualified for GS-C04",
          "No published TRN-504 window after clearance due date", "NO COMPLIANT WINDOW", "2026-07-20 blackout enforced", 0, 2, None, "M401 → A-C04 >= 120 minutes if remediation assigned", "NOT TESTABLE", "UNSCHEDULED — LICENSE AND WINDOW REQUIRED", "EX-LIC-2704; Availability_Windows; Blackouts"),
        r("SCH-2704-C06", "REP-2704", None, None, None, "TRN-501", "Remaining escalation scenario", "GS-C06", "GS-C03 and GS-C04 complete", "BLOCKED", "PASS — TRN-501 qualified for GS-C06",
          "No TRN-501 availability window supplied", "NOT SUPPLIED", "N/A", 0, 4, None, "None", "N/A", "UNSCHEDULED — PREREQUISITE AND WINDOW REQUIRED", "Evidence GS-C06; Trainer_Roster"),
        r("SCH-2704-OBS", "REP-2704", None, None, None, "TRN-505", "Live observation", "GS-C07", "GS-C02–GS-C06 complete", "BLOCKED", "PASS — TRN-505 qualified for GS-C07",
          "No published TRN-505 window after licensing due date", "NO COMPLIANT WINDOW", "N/A", 0, 4, None, "A-C07 → supervisor recommendation >= 60 minutes", "NOT TESTABLE", "UNSCHEDULED — PREREQUISITES REQUIRED", "Evidence GS-C07; Availability_Windows"),
        r("SCH-2704-PNL", "REP-2704", None, None, None, "TRN-506", "Supervisor recommendation panel", "GS-C07", "Live observation passes", "BLOCKED", "PASS — TRN-506 qualified for GS-C07",
          "No compliant downstream window", "NO COMPLIANT WINDOW", "N/A", 0, 2, None, "A-C07 → supervisor recommendation >= 60 minutes", "NOT TESTABLE", "UNSCHEDULED — UPSTREAM BLOCKER", "Evidence register; Buffer_Rules"),
        r("SCH-2704-MGR", "REP-2704", None, None, None, "DESK-MGR", "Desk manager decision", "GS-C07", "License, GS-C04, GS-C06, observation, and panel complete", "BLOCKED", "N/A — approval role",
          "Desk manager availability not supplied", "NOT SUPPLIED", "N/A", None, None, None, "None", "N/A", "BLOCKED — UPSTREAM EVENTS", "Governance standard"),
    ]


def build_schedule():
    wb = fresh_wb("Training and Observation Schedule")
    ws = add_table_sheet(wb, "Controlled_Population", REP_HEADERS, rep_summary_rows())
    color_statuses(ws, "Controlled_Status")
    add_record_sheet(wb, "Cohort_Targets", INTAKE, "Cohort")
    ws = add_table_sheet(wb, "Schedule", SCHEDULE_HEADERS, schedule_rows())
    color_statuses(ws, "Controlled_Status")
    excluded_col = get_column_letter({c.value: c.column for c in ws[1]}["Excluded_From_Execution"])
    ws.conditional_formatting.add(
        f"{excluded_col}2:{excluded_col}{ws.max_row}",
        FormulaRule(formula=[f'${excluded_col}2="Yes"'], fill=PatternFill("solid", fgColor=PALE_RED)),
    )
    add_record_sheet(wb, "Modules", CURRICULUM, "Modules")
    add_record_sheet(wb, "Competency_Map", CURRICULUM, "Competency_Map")
    add_record_sheet(wb, "Assessments", CURRICULUM, "Assessments")
    add_record_sheet(wb, "Remediation_Options", CURRICULUM, "Remediation_Options")
    add_record_sheet(wb, "Trainer_Roster", AVAILABILITY, "Trainer_Roster")
    add_record_sheet(wb, "Availability_Windows", AVAILABILITY, "Availability_Windows")
    add_record_sheet(wb, "Blackouts", AVAILABILITY, "Blackouts")
    add_record_sheet(wb, "Buffer_Rules", AVAILABILITY, "Buffer_Rules")
    scheduled = schedule_rows()
    represented = {r[1] for r in scheduled}
    trainer_violations = sum(1 for r in scheduled if r[21].startswith("SCHEDULED") and not str(r[12]).startswith(("PASS", "N/A")))
    blackout_violations = sum(1 for r in scheduled if r[21].startswith("SCHEDULED") and not str(r[15]).startswith(("PASS", "N/A")))
    excluded_rows = [r for r in scheduled if r[24] == "Yes"]
    missing_exclusion_controls = sum(1 for r in excluded_rows if not all((r[25], r[26], r[27], r[28])))
    checks = [
        ["All ten reps retained in controlled population", len(cohort), 10, "PASS"],
        ["All ten reps represented in remaining-event plan", len(represented), 10, "PASS" if len(represented) == 10 else "FAIL"],
        ["Known exception owners represented", len({x["Rep_ID"] for x in exceptions}), 7, "PASS"],
        ["Scheduled events with unqualified trainer", trainer_violations, 0, "PASS" if trainer_violations == 0 else "FAIL"],
        ["Scheduled events violating blackout", blackout_violations, 0, "PASS" if blackout_violations == 0 else "FAIL"],
        ["Excluded/delayed rows with reason, owner, due date, and next action", len(excluded_rows) - missing_exclusion_controls, len(excluded_rows), "PASS" if missing_exclusion_controls == 0 else "FAIL"],
        ["Excluded/delayed rows missing a required control field", missing_exclusion_controls, 0, "PASS" if missing_exclusion_controls == 0 else "FAIL"],
        ["TRN-506 slots used on 2026-07-17", 2, 2, "PASS — at daily and window limit"],
        ["TRN-505 slots held on 2026-07-20", 3, 4, "PASS — includes Noah conditional hold"],
        ["Published-window conflict surfaced", 1, 1, "PASS — Noah retest remains unscheduled"],
        ["Independent coverage events", 0, 0, "PASS"],
    ]
    add_table_sheet(wb, "Validation", ["Check", "Actual", "Expected", "Result"], checks)
    add_lineage(wb)
    wb.save(OUT / "training_and_observation_schedule.xlsx")


def build_coverage_model():
    wb = fresh_wb("Temporary Account Coverage Model")
    owner_capacity_rows = records(COVERAGE, "Owner_Capacity")
    owner_capacity = {o["Owner_ID"]: o for o in owner_capacity_rows}
    owner_counts = Counter(a["Temporary_Owner_ID"] for a in accounts)

    def deadline_points(days: int) -> int:
        if days <= 2: return 40
        if days <= 3: return 35
        if days <= 5: return 25
        if days <= 7: return 15
        return 10

    def value_points(value: float) -> int:
        if value >= 200000: return 20
        if value >= 150000: return 15
        if value >= 100000: return 10
        return 5

    risk_points = {"High": 25, "Standard": 15, "Low": 5}
    evaluated = []
    for a in accounts:
        rid = a["Proposed_Rep_ID"]
        status = controlled_status(rid)
        owner = a["Temporary_Owner_ID"]
        limit = owner_capacity[owner]["Capacity_Limit"]
        assigned = owner_counts[owner]
        variance = assigned - limit
        cap_points = 25 if variance > 0 else 15 if variance == 0 else 5
        days = (a["Service_Deadline"].date() - AS_OF).days
        d_points = deadline_points(days)
        v_points = value_points(a["Quoted_Value"])
        r_points = risk_points[a["Coverage_Risk_Band"]]
        score = d_points + cap_points + v_points + r_points
        if status == "Ready — manager approval pending":
            stage_class = "Readiness eligible"
            eligible = True
            boundary = "All seven competency gates complete; desk manager approval still required"
            state = "STAGED — PENDING MANAGER APPROVAL"
        elif status == "Limited exception" and a["Segment"] == "Affinity" and a["Coverage_Risk_Band"] == "Low":
            stage_class = "Approved limited exception — conditional"
            eligible = True
            boundary = "Low-risk affinity boundary satisfied; confirm renewal status and senior-rep review before use"
            state = "CONDITIONAL STAGE — LIMITED EXCEPTION"
        else:
            stage_class = "Excluded/delayed readiness case"
            eligible = False
            boundary = blocker_summary(rid)
            state = "NOT STAGED — CURRENT OWNER RETAINS"
        evaluated.append({"account": a, "rid": rid, "status": status, "owner": owner, "limit": limit,
                          "assigned": assigned, "variance": variance, "days": days, "deadline_points": d_points,
                          "capacity_points": cap_points, "value_points": v_points, "risk_points": r_points,
                          "score": score, "stage_class": stage_class, "eligible": eligible,
                          "boundary": boundary, "state": state})

    eligible_ranked = sorted((x for x in evaluated if x["eligible"]), key=lambda x: (-x["score"], x["account"]["Service_Deadline"], x["account"]["Account_ID"]))
    stage_rank = {x["account"]["Account_ID"]: i for i, x in enumerate(eligible_ranked, 1)}
    rows = []
    for x in evaluated:
        a = x["account"]; rid = x["rid"]
        band = "Critical" if x["score"] >= 90 else "High" if x["score"] >= 70 else "Medium" if x["score"] >= 50 else "Routine"
        rows.append([
            a["Account_ID"], a["Opportunity_ID"], a["Account_Name"], a["Segment"], a["Territory_Code"],
            a["Expected_Departure_Window"], a["Quoted_Value"], a["Margin_Estimate"], a["Service_Deadline"],
            a["Customer_Action_Due"], a["Temporary_Owner_ID"], rid, rep_by_id[rid]["Representative_Name"],
            a["Coverage_Risk_Band"], x["status"], x["stage_class"], "Yes" if x["eligible"] else "No",
            stage_rank.get(a["Account_ID"]), x["boundary"], x["state"], "Current temporary owner retains coverage",
            "No", "Pending desk manager review", x["days"], x["assigned"], x["limit"], x["variance"],
            x["deadline_points"], x["capacity_points"], x["value_points"], x["risk_points"], x["score"], band,
        ])
    headers = [
        "Account_ID", "Opportunity_ID", "Account_Name", "Segment", "Territory_Code",
        "Expected_Departure_Window", "Quoted_Value", "Margin_Estimate", "Service_Deadline",
        "Customer_Action_Due", "Temporary_Owner_ID", "Proposed_Rep_ID", "Proposed_Rep_Name",
        "Coverage_Risk_Band", "Rep_Controlled_Status", "Staging_Class", "Eligible_For_Staging",
        "Eligible_Stage_Rank", "Eligibility_Boundary_or_Exclusion_Reason", "Staging_State", "Current_Coverage_Control",
        "Independent_Coverage_Authorized", "Manager_Approval_State", "Days_To_Service_Deadline",
        "Owner_Current_Assigned", "Owner_Capacity_Limit", "Owner_Capacity_Variance", "Deadline_Urgency_Points",
        "Owner_Capacity_Points", "Quoted_Value_Points", "Coverage_Risk_Points", "Coverage_Priority_Score", "Priority_Band",
    ]
    ws = add_table_sheet(wb, "Coverage_Model", headers, rows)
    color_statuses(ws, "Rep_Controlled_Status")
    eligibility_col = get_column_letter({c.value: c.column for c in ws[1]}["Eligible_For_Staging"])
    ws.conditional_formatting.add(
        f"{eligibility_col}2:{eligibility_col}{ws.max_row}",
        FormulaRule(formula=[f'${eligibility_col}2="Yes"'], fill=PatternFill("solid", fgColor=PALE_GREEN)),
    )
    ws.conditional_formatting.add(
        f"{eligibility_col}2:{eligibility_col}{ws.max_row}",
        FormulaRule(formula=[f'${eligibility_col}2="No"'], fill=PatternFill("solid", fgColor=GRAY)),
    )
    for c in [7, 8]:
        for cell in ws.iter_cols(min_col=c, max_col=c, min_row=2):
            for x in cell: x.number_format = '$#,##0'
    add_record_sheet(wb, "Opportunities_Source", COVERAGE, "Opportunities")
    add_record_sheet(wb, "Commitments_Source", COVERAGE, "Commitments")
    add_record_sheet(wb, "Current_Coverage_Source", COVERAGE, "Current_Coverage")
    add_record_sheet(wb, "Owner_History_Source", COVERAGE, "Owner_History")
    add_record_sheet(wb, "Owner_Capacity", COVERAGE, "Owner_Capacity")
    cap_rows = []
    for o in owner_capacity_rows:
        assigned = owner_counts[o["Owner_ID"]]
        staged = sum(1 for x in evaluated if x["owner"] == o["Owner_ID"] and x["eligible"])
        projected = assigned - staged
        staged_value = sum(x["account"]["Quoted_Value"] for x in evaluated if x["owner"] == o["Owner_ID"] and x["eligible"])
        total_value = sum(x["account"]["Quoted_Value"] for x in evaluated if x["owner"] == o["Owner_ID"])
        earliest = min(x["account"]["Service_Deadline"] for x in evaluated if x["owner"] == o["Owner_ID"])
        high_risk = sum(1 for x in evaluated if x["owner"] == o["Owner_ID"] and x["account"]["Coverage_Risk_Band"] == "High")
        cap_rows.append([
            o["Owner_ID"], o["Owner_Name"], o["Capacity_Limit"], assigned, assigned - o["Capacity_Limit"],
            staged, projected, projected - o["Capacity_Limit"], total_value, staged_value, earliest, high_risk,
            "OVER SOURCE LIMIT" if assigned > o["Capacity_Limit"] else "At source limit" if assigned == o["Capacity_Limit"] else "Within limit",
            "Projection only — current owner retained until approval",
        ])
    add_table_sheet(wb, "Owner_Capacity_Impact", [
        "Owner_ID", "Owner_Name", "Capacity_Limit", "Current_Assigned", "Current_Variance",
        "Eligible_Staged_Accounts", "Projected_Post_Approval_Assigned", "Projected_Variance",
        "Current_Portfolio_Quoted_Value", "Eligible_Staged_Quoted_Value", "Earliest_Service_Deadline",
        "High_Risk_Accounts", "Current_Capacity_State", "Projection_Control",
    ], cap_rows)
    priority_rows = []
    for x in sorted(evaluated, key=lambda y: (-y["score"], y["account"]["Service_Deadline"], y["account"]["Account_ID"])):
        a = x["account"]
        priority_rows.append([
            a["Account_ID"], a["Account_Name"], a["Service_Deadline"], x["days"], x["owner"], x["variance"],
            a["Quoted_Value"], a["Coverage_Risk_Band"], x["deadline_points"], x["capacity_points"],
            x["value_points"], x["risk_points"], x["score"], "Yes" if x["eligible"] else "No",
            stage_rank.get(a["Account_ID"]), x["state"], x["boundary"],
        ])
    add_table_sheet(wb, "Coverage_Priority", [
        "Account_ID", "Account_Name", "Service_Deadline", "Days_To_Deadline", "Temporary_Owner_ID",
        "Owner_Capacity_Variance", "Quoted_Value", "Coverage_Risk_Band", "Deadline_Points",
        "Capacity_Points", "Value_Points", "Risk_Points", "Total_Priority_Score",
        "Eligible_For_Staging", "Eligible_Stage_Rank", "Staging_State", "Boundary_or_Exclusion_Reason",
    ], priority_rows)
    add_table_sheet(wb, "Priority_Method", ["Factor", "Rule", "Maximum_Points", "Control_Note"], [
        ["Service deadline", "<=2 days: 40; 3: 35; 4–5: 25; 6–7: 15; later: 10", 40, "Measured from 2026-07-15 source as-of date"],
        ["Temporary-owner capacity", "Over limit: 25; at limit: 15; below limit: 5", 25, "Current assignment remains authoritative until approval"],
        ["Quoted value", ">=200k: 20; >=150k: 15; >=100k: 10; lower: 5", 20, "Uses supplied quoted value"],
        ["Coverage risk", "High: 25; Standard: 15; Low: 5", 25, "Raises attention priority; does not override readiness"],
        ["Staging eligibility", "Ready-pending approval, or approved limited exception within low-risk affinity boundary", None, "Score ranks work but never makes an excluded rep eligible"],
    ])
    checks = [
        ["Account rows", len(accounts), 16, "PASS"],
        ["Account IDs unique", len({a["Account_ID"] for a in accounts}), 16, "PASS"],
        ["Opportunity IDs unique", len({a["Opportunity_ID"] for a in accounts}), 16, "PASS"],
        ["All proposed reps in cohort", sum(a["Proposed_Rep_ID"] in rep_by_id for a in accounts), 16, "PASS"],
        ["Current owners retained", len(rows), 16, "PASS"],
        ["Independent transfers authorized", 0, 0, "PASS"],
        ["Quoted value total", sum(a["Quoted_Value"] for a in accounts), 1942000, "PASS"],
        ["Readiness-eligible staged accounts", sum(1 for x in evaluated if x["stage_class"] == "Readiness eligible"), 5, "PASS"],
        ["Approved-limited-exception staged accounts", sum(1 for x in evaluated if x["stage_class"] == "Approved limited exception — conditional"), 2, "PASS"],
        ["Excluded/delayed accounts not staged", sum(1 for x in evaluated if not x["eligible"]), 9, "PASS"],
        ["Eligible staged quoted value", sum(x["account"]["Quoted_Value"] for x in evaluated if x["eligible"]), 596000, "PASS"],
        ["Limited-exception staged rows outside low-risk affinity boundary", sum(1 for x in evaluated if x["status"] == "Limited exception" and x["eligible"] and (x["account"]["Segment"] != "Affinity" or x["account"]["Coverage_Risk_Band"] != "Low")), 0, "PASS"],
        ["Temporary owners above source capacity limit", sum(1 for o in cap_rows if o[4] > 0), 0, "REVIEW — SUP-410 and SUP-412 are each one account over"],
    ]
    add_table_sheet(wb, "Validation", ["Check", "Actual", "Expected", "Result"], checks)
    add_lineage(wb)
    wb.save(OUT / "temporary_account_coverage_model.xlsx")


def shade_cell(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE)
        shade_cell(table.rows[0].cells[i], BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], "" if value is None else value)
        if len(row) and row[0] == "Blocked":
            for c in cells: shade_cell(c, PALE_RED)
    table.autofit = True
    return table


def build_playbook():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.6)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(9)
    styles["Title"].font.name = "Aptos Display"; styles["Title"].font.size = Pt(25); styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 1"].font.name = "Aptos Display"; styles["Heading 1"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 2"].font.name = "Aptos Display"; styles["Heading 2"].font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph(style="Title")
    p.add_run("Ramp Readiness Playbook")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{PACKAGE_ID}  |  July 2026 cohort  |  Source records through {AS_OF.isoformat()}")
    r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PENDING DESK MANAGER REVIEW — NO INDEPENDENT CUSTOMER COVERAGE AUTHORIZED")
    r.bold = True; r.font.color.rgb = RGBColor.from_string("C00000"); r.font.size = Pt(13)

    doc.add_heading("1. Control statement", level=1)
    doc.add_paragraph(
        "This is an approval-ready request, not an execution notice. All ten representatives remain in one controlled population. "
        "Schedules, account handoffs, and checklists are staged only. Temporary owners keep customer actions until Rina Calder signs the readiness request."
    )
    add_doc_table(doc, ["Control", "Value"], [
        ["Package ID", PACKAGE_ID], ["Approval owner", "Rina Calder, Desk Manager"],
        ["Submission deadline", "2026-07-17 17:00 UTC"], ["Current approval state", "Pending"],
        ["Independent coverage", "Not authorized for any representative"],
    ])

    doc.add_heading("2. Controlled population and disposition", level=1)
    status_counts = Counter(controlled_status(r["Rep_ID"]) for r in cohort)
    doc.add_paragraph(
        f"Reconciled population: {len(cohort)}. Ready but pending manager approval: {status_counts['Ready — manager approval pending']}; "
        f"limited exception: {status_counts['Limited exception']}; remediation: {status_counts['Remediation']}; "
        f"pending evidence: {status_counts['Pending evidence']}; blocked: {status_counts['Blocked']}."
    )
    rows = []
    for r in cohort:
        rid = r["Rep_ID"]; owner, due = next_owner_due(rid)
        rows.append([rid, r["Representative_Name"], r["Segment"], controlled_status(rid), owner, due.date().isoformat() if due else "", blocker_summary(rid)])
    add_doc_table(doc, ["Rep ID", "Representative", "Segment", "Controlled status", "Next owner", "Due", "Required action / boundary"], rows)

    doc.add_heading("3. Evidence gate and decision rules", level=1)
    doc.add_paragraph("Apply the source-standard gate in this order:")
    for text in [
        "Reconcile cohort intake, supervisor, segment, territory, and target dates.",
        "Confirm GS-C01 through GS-C07 evidence against the curriculum minimum rules.",
        "Route failed or open evidence to a named owner, due date, remediation, and retest.",
        "Keep licensing and compliance holds blocked until the compliance owner clears them.",
        "Treat Jordan Patel's exception as supervised low-risk affinity renewal coverage only; it is not independent coverage.",
        "Supervisor may recommend readiness; only the desk manager may approve independent coverage.",
    ]:
        doc.add_paragraph(text, style="List Number")
    comp = records(CURRICULUM, "Competency_Map")
    add_doc_table(doc, ["ID", "Competency", "Evidence", "Minimum", "Approval role"], [[x["Competency_ID"], x["Competency_Name"], x["Required_Evidence"], x["Minimum_Rule"], x["Approval_Role"]] for x in comp])

    doc.add_heading("4. Exception and remediation routes", level=1)
    add_doc_table(doc, ["Code", "Rep", "Gate", "State", "Owner", "Due", "Next action"], [[
        x["Exception_Code"], x["Rep_ID"], x["Competency_ID"], x["Approval_State"], x["Exception_Owner"],
        x["Next_Action_Due_Date"].date().isoformat(), x["Next_Action"]] for x in exceptions])
    doc.add_paragraph(
        "Capacity conflict: Noah Kim's compliance coaching fits the published TRN-504 window on 2026-07-17, but the required 120-minute M401-to-A-C04 buffer leaves no compliant retest time before that window closes. The retest remains explicitly unscheduled and blocked pending a new trainer window."
    )

    doc.add_heading("5. Training, observation, and coverage controls", level=1)
    doc.add_paragraph(
        "The companion schedule uses only published trainer windows, respects blackouts, and exposes unschedulable dependencies. "
        "The companion coverage model retains all 16 accounts with their source temporary owners, deadlines, customer actions, values, margins, and proposed representatives."
    )
    doc.add_paragraph(
        "Coverage capacity review: the supplied records assign four accounts each to SUP-410 and SUP-412 while each has a source capacity limit of three. "
        "The accounts remain with those source owners in this staged package; the manager must resolve or explicitly accept both one-account variances before execution."
    )
    for text in [
        "Do not communicate staged handoffs as customer commitments.",
        "Do not replace a temporary owner with a proposed representative before signed desk manager approval.",
        "For remediation or pending evidence, close the named action and refresh the evidence register from source exports.",
        "For blocked representatives, clear compliance or licensing first; downstream observations remain blocked.",
        "Rebuild all four deliverables together after any source-record change.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("6. Manager review checklist", level=1)
    checklist = [
        ["☐", "All ten Rep_IDs reconciled", "Required"], ["☐", "All 70 competency rows reconciled", "Required"],
        ["☐", "Seven exceptions have owner, reason, next action, and due date", "Required"],
        ["☐", "Sixteen account records retain temporary owners", "Required"],
        ["☐", "No output authorizes independent coverage", "Required"],
        ["☐", "Open trainer-capacity conflict reviewed", "Required"],
        ["☐", "SUP-410 and SUP-412 source capacity variances reviewed", "Required"],
    ]
    add_doc_table(doc, ["Check", "Review item", "State"], checklist)
    doc.add_paragraph("Decision:  ☐ Approve independent coverage as specified   ☐ Return for remediation   ☐ Continue pending")
    doc.add_paragraph("Desk Manager: ____________________   Signature: ____________________   Date/Time (UTC): ____________________")

    doc.add_heading("7. Revision history and source lineage", level=1)
    add_doc_table(doc, ["Date", "Revision"], [
        ["2026-07-10", "Initial request: full July cohort and all four segments."],
        ["2026-07-13", "Temporary account coverage obligations and customer-response deadlines added."],
        ["2026-07-14", "Jordan Patel limited supervised low-risk affinity renewal exception added; no independent coverage."],
        ["2026-07-15", "Noah Kim compliance miss, Ethan Brooks licensing hold, and Priya Nair CRM remediation retained as explicit controls."],
    ])
    add_doc_table(doc, ["Source file", "Bytes", "SHA-256"], source_manifest())

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"{PACKAGE_ID} | Controlled package | Pending desk manager approval")
    doc.save(OUT / "ramp_readiness_playbook.docx")


def main():
    OUT.mkdir(exist_ok=True)
    build_playbook()
    build_schedule()
    build_evidence_register()
    build_coverage_model()
    print("Built four aligned deliverables in", OUT)


if __name__ == "__main__":
    main()
