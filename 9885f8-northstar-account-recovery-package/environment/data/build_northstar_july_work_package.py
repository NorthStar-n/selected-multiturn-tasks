from __future__ import annotations

import csv
import re
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path

from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import CellIsRule, FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document


ROOT = Path(__file__).resolve().parent
DATA = ROOT / "data"
OUTPUT = ROOT / "output"
OUTPUT_FILE = OUTPUT / "northstar_july_recovery_sprint_work_package.xlsx"

SUPERVISOR = "Dana Morales"
CHECKPOINT = "16:00 Eastern"
SPRINT_START = date(2026, 7, 17)
SPRINT_END = date(2026, 7, 24)

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
TEAL = "0F6B78"
LIGHT_TEAL = "DDEBF7"
GREEN = "548235"
LIGHT_GREEN = "E2F0D9"
AMBER = "BF9000"
LIGHT_AMBER = "FFF2CC"
RED = "C00000"
LIGHT_RED = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "E7E6E6"
WHITE = "FFFFFF"

THIN_GRAY = Side(style="thin", color="B7B7B7")
GRID = Border(left=THIN_GRAY, right=THIN_GRAY, top=THIN_GRAY, bottom=THIN_GRAY)


def read_sheet(path: Path, sheet: str) -> list[dict]:
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb[sheet]
    rows = ws.iter_rows(values_only=True)
    headers = [str(v) for v in next(rows)]
    return [dict(zip(headers, row)) for row in rows if any(v is not None for v in row)]


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def add_title(ws, title: str, subtitle: str, width: int) -> None:
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=width)
    ws.cell(1, 1, title)
    ws.cell(1, 1).font = Font(size=18, bold=True, color=WHITE)
    ws.cell(1, 1).fill = PatternFill("solid", fgColor=NAVY)
    ws.cell(1, 1).alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 28
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=width)
    ws.cell(2, 1, subtitle)
    ws.cell(2, 1).font = Font(size=10, italic=True, color=GRAY)
    ws.cell(2, 1).alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[2].height = 30


def style_header(ws, row: int, start: int, end: int, color: str = BLUE) -> None:
    for col in range(start, end + 1):
        cell = ws.cell(row, col)
        cell.font = Font(bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=color)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = GRID
    ws.row_dimensions[row].height = 32


def style_body(ws, start_row: int, end_row: int, end_col: int) -> None:
    for row in ws.iter_rows(min_row=start_row, max_row=end_row, min_col=1, max_col=end_col):
        for cell in row:
            cell.border = GRID
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for r in range(start_row, end_row + 1):
        ws.row_dimensions[r].height = 48


def add_table(ws, ref: str, name: str, style: str = "TableStyleMedium2") -> None:
    table = Table(displayName=name, ref=ref)
    table.tableStyleInfo = TableStyleInfo(
        name=style, showFirstColumn=False, showLastColumn=False,
        showRowStripes=True, showColumnStripes=False
    )
    ws.add_table(table)


def set_widths(ws, widths: dict[str, float]) -> None:
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def as_date(value):
    if isinstance(value, datetime):
        return value.date()
    return value


watchlist_path = DATA / "revenue_ops_review" / "escalation_watchlist_review_2026-07-14.xlsx"
crm_path = DATA / "crm_health" / "escalated_account_health_export_2026-07-13.xlsx"
contacts_path = DATA / "contacts" / "advertiser_contact_and_authority_directory_2026-07-10.xlsx"
cases_path = DATA / "ad_ops_cases" / "open_ad_ops_case_queue_2026-07-13.csv"
activity_path = DATA / "rep_activity" / "rep_touchpoints_and_commitments_2026-04-01_to_2026-07-14.csv"
authority_path = DATA / "commercial_options" / "local_recovery_offer_authority_q3_2026.docx"
approval_path = DATA / "approvals" / "recovery_sprint_scope_approval_2026-07-15.eml"

watchlist = read_sheet(watchlist_path, "Watchlist Data")
accounts = {r["Account ID"]: r for r in read_sheet(crm_path, "Accounts")}
opportunities = {r["Account ID"]: r for r in read_sheet(crm_path, "Opportunities")}
contacts = read_sheet(contacts_path, "Contacts")
cases = read_csv(cases_path)
activities = read_csv(activity_path)


def approval_ids_between(text: str, start: str, end: str | None) -> set[str]:
    section = text.split(start, 1)[1]
    if end:
        section = section.split(end, 1)[0]
    return set(re.findall(r"A-\d{4}", section))


approval_text = approval_path.read_text(encoding="utf-8")
approved_scope_ids = approval_ids_between(
    approval_text, "Approved for local execution:", "Controlled exceptions"
)
controlled_scope_ids = approval_ids_between(
    approval_text, "Controlled exceptions that must remain visible", "Excluded or hold records"
)
excluded_scope_ids = approval_ids_between(
    approval_text, "Excluded or hold records for this sprint:", "Approved offer bands"
)
if (len(approved_scope_ids), len(controlled_scope_ids), len(excluded_scope_ids)) != (12, 4, 2):
    raise ValueError("July 15 approval scope did not reconcile to 12 approved, 4 controlled, and 2 excluded/hold accounts.")

authority_doc = Document(authority_path)
authority_source: dict[str, dict[str, str]] = {}
for table in authority_doc.tables:
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    if headers and headers[0] == "Offer Code":
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            authority_source[values[0]] = dict(zip(headers, values))
        break

if set(authority_source) != {"MG-15", "CR-10K", "RH-60", "SW-PKG", "RPT-COMP", "MGR-REQ"}:
    raise ValueError("Offer authority matrix did not reconcile to the expected six source codes.")

contacts_by_account: dict[str, list[dict]] = defaultdict(list)
for record in contacts:
    contacts_by_account[record["Account ID"]].append(record)

cases_by_account: dict[str, list[dict]] = defaultdict(list)
for record in cases:
    cases_by_account[record["account_id"]].append(record)

activities_by_account: dict[str, list[dict]] = defaultdict(list)
for record in activities:
    activities_by_account[record["account_id"]].append(record)

watch_by_id = {r["Account ID"]: r for r in watchlist}

# The schedule is deliberately gated: a planned date is not permission to contact.
execution_plan = {
    "A-1005": (1, date(2026, 7, 20), "Audience rebuild proof and package-swap feasibility confirmed by Product Support / Ad Ops.",
               "Describe the equivalent-value package swap only after feasibility is documented; do not guarantee future performance."),
    "A-1001": (2, date(2026, 7, 17), "Traffic pacing confirmation and makegood inventory validation attached to CASE-7201.",
               "Describe the validated makegood inventory and revised delivery plan; do not promise unvalidated impressions."),
    "A-1012": (3, date(2026, 7, 20), "Makegood placement validated and invoice PO correction route documented.",
               "Describe the validated makegood placement and invoice-correction process; do not combine this with an unapproved credit."),
    "A-1007": (4, date(2026, 7, 20), "Makegood placement and revised delivery dates confirmed by Ad Ops.",
               "Describe the validated underdelivery remedy and dates; do not guarantee future delivery."),
    "A-1003": (5, date(2026, 7, 17), "Verified credit memo up to $9,500 attached; any rate-hold request removed or separately escalated.",
               "Discuss only the verified invoice credit under CR-10K. Do not offer a local rate hold because Key Growth is not RH-60 eligible."),
    "A-1009": (6, date(2026, 7, 17), "Creative approval proof and equivalent package-swap feasibility confirmed by Ad Ops.",
               "Describe the approved equivalent-value swap and confirmed timeline; do not promise launch timing before validation."),
    "A-1002": (7, date(2026, 7, 21), "Corrected reach report and customer-facing reconciliation summary attached.",
               "Provide the reconciled reporting package; make no commercial or performance promise."),
    "A-1004": (8, date(2026, 7, 17), "Creative retagging confirmation and package-swap feasibility documented by Ad Ops.",
               "Describe the approved equivalent-value swap after trafficking confirmation; do not promise timing before validation."),
    "A-1011": (9, date(2026, 7, 21), "Corrected lift model and customer-facing summary attached.",
               "Provide the corrected reporting package; make no commercial or performance promise."),
    "A-1006": (10, date(2026, 7, 20), "Verified credit memo up to $4,800 and corrected bill date attached.",
               "Discuss only the verified invoice correction under CR-10K; do not exceed $4,800."),
    "A-1008": (11, date(2026, 7, 21), "Corrected district reach report and renewal summary attached.",
               "Provide the reconciled district report; make no commercial or performance promise."),
    "A-1010": (12, date(2026, 7, 22), "60-day CPM hold language documented with no additional credit.",
               "Offer only the 60-day current-CPM hold under RH-60; do not add a credit."),
}

offer_caps = {
    "MG-15": "$25,000 / up to 15% of affected booked impressions",
    "CR-10K": "$10,000 verified invoice credit",
    "RH-60": "No credit; current CPM held for 60 days",
    "SW-PKG": "$18,000 equivalent-value package swap",
    "RPT-COMP": "No credit; reporting reconciliation package",
}

exception_owner = {
    "A-1013": ("Supervisor / contact-route approval", date(2026, 7, 20), "Obtain permission for the final approver route; operational contact may receive issue evidence only."),
    "A-1014": ("Sales manager", date(2026, 7, 17), "Decide the $42,000 over-limit credit route; no local credit promise."),
    "A-1015": ("Product Support / Ad Ops", date(2026, 7, 22), "Clear the feed blocker and provide feasibility proof before any package-swap offer."),
    "A-1017": ("Sales manager and Legal", date(2026, 7, 24), "Draft the data-use response through the manager/legal route; no rep commitment."),
}

# Cases that control the local action's expected evidence date. Other open cases remain visible,
# but do not postpone an unrelated authorized action (notably A-1003's unsupported rate request).
action_case_ids = {
    "A-1001": {"CASE-7201"}, "A-1002": {"CASE-7203"}, "A-1003": {"CASE-7204"},
    "A-1004": {"CASE-7206"}, "A-1005": {"CASE-7207", "CASE-7208"}, "A-1006": {"CASE-7209"},
    "A-1007": {"CASE-7210"}, "A-1008": {"CASE-7211"}, "A-1009": {"CASE-7212"},
    "A-1010": {"CASE-7213"}, "A-1011": {"CASE-7214"}, "A-1012": {"CASE-7215", "CASE-7216"},
    "A-1013": {"CASE-7217"}, "A-1014": {"CASE-7218"}, "A-1015": {"CASE-7219"},
    "A-1017": {"CASE-7221", "CASE-7222"}, "A-1018": {"CASE-7223"},
}


def primary_contact(account_id: str) -> dict:
    records = contacts_by_account[account_id]
    approved_final = [r for r in records if r["Decision Authority"] == "Yes" and r["Outreach Permission"] == "Approved"]
    if approved_final:
        return approved_final[0]
    approved_ops = [r for r in records if r["Outreach Permission"] in ("Approved", "Operations Only")]
    return approved_ops[0] if approved_ops else records[0]


def current_commitment(account_id: str) -> tuple[str, str]:
    customer_records = [r for r in activities_by_account[account_id] if r["activity_type"] in ("Email", "Customer call")]
    record = sorted(customer_records, key=lambda r: r["activity_date"])[-1]
    return record["internal_commitment"], record["commitment_due_date"]


def case_summary(account_id: str) -> tuple[str, str, str]:
    records = [r for r in cases_by_account[account_id] if r["issue_status"] != "Closed"]
    case_ids = ", ".join(r["case_id"] for r in records)
    owners = ", ".join(dict.fromkeys(r["current_owner"] for r in records))
    controlling = [r for r in records if r["case_id"] in action_case_ids.get(account_id, set())]
    expected = max((r["expected_resolution_date"] for r in controlling), default="")
    return case_ids, owners, expected


def july_scope(account_id: str) -> str:
    if account_id in approved_scope_ids:
        return "Approved local execution"
    if account_id in controlled_scope_ids:
        return "Controlled exception"
    if account_id in excluded_scope_ids:
        return "Excluded / hold"
    return "Not named"


def package_lane(account_id: str) -> str:
    if account_id in approved_scope_ids:
        return "WORK NOW — CLEAR / PREPARE"
    if account_id in controlled_scope_ids:
        return "CONTROLLED — NO LOCAL COMMITMENT"
    return "EXCLUDED / HOLD"


def contact_route_summary(account_id: str) -> tuple[str, str]:
    finals = [r for r in contacts_by_account[account_id] if r["Decision Authority"] == "Yes"]
    final = finals[0]
    route = f'{final["Contact Name"]}: {final["Outreach Permission"]}'
    if account_id in approved_scope_ids:
        check = "Pass" if final["Outreach Permission"] == "Approved" else "Mismatch"
    elif account_id == "A-1013":
        check = "Controlled — final approver permission pending"
    elif account_id == "A-1017":
        check = "Controlled — executive do-not-contact / manager-legal route"
    elif account_id == "A-1018":
        check = "Hold — customer requested September review"
    else:
        check = "Consistent with hold/exception lane"
    return route, check


def parse_money(value: str) -> int:
    digits = re.sub(r"[^0-9]", "", value or "")
    return int(digits) if digits else 0


def offer_reconciliation(record: dict) -> str:
    code = record["Approved Offer"]
    aid = record["Account ID"]
    if not code:
        return "Not applicable — no local offer"
    if code == "MGR-REQ":
        return "Controlled manager route — not a local offer"
    source = authority_source[code]
    tiers = {v.strip() for v in source["Eligible Tiers"].split(",")}
    issues = {v.strip() for v in source["Issue Basis"].split(";")}
    cap = parse_money(source["Max Value"])
    tier_ok = record["Account Tier"] in tiers
    issue_ok = record["Primary Issue"] in issues
    value_ok = float(record["Requested Recovery Value"] or 0) <= cap if cap else float(record["Requested Recovery Value"] or 0) == 0
    base = "Pass" if tier_ok and issue_ok and value_ok else "Mismatch"
    if aid == "A-1013":
        return f"{base} offer fit; controlled by contact permission"
    if aid == "A-1015":
        return f"{base} offer fit; controlled by unresolved ops blocker"
    if aid == "A-1003":
        return f"{base} for CR-10K only; RH-60 excluded for Key Growth"
    return base


def ops_reconciliation(record: dict) -> tuple[str, str, str, str]:
    aid = record["Account ID"]
    open_records = [r for r in cases_by_account[aid] if r["issue_status"] != "Closed"]
    count_ok = len(open_records) == int(record["Open Cases"] or 0)
    actual_max = max((int(r["severity_score"]) for r in open_records), default=0)
    severity_ok = actual_max == int(record["Max Severity"] or 0)
    actual_avg = sum(int(r["issue_age_days"]) for r in open_records) / len(open_records) if open_records else 0
    age_ok = abs(actual_avg - float(record["Avg Issue Age"] or 0)) < 0.01
    statuses = "; ".join(f'{r["case_id"]} {r["issue_status"]}' for r in open_records) or "No open cases"
    dependencies = "; ".join(f'{r["case_id"]}: {r["open_dependency"]}' for r in open_records) or "None"
    return (
        "Pass" if count_ok else f"Mismatch: ops {len(open_records)} / watchlist {record['Open Cases']}",
        "Pass" if severity_ok else f"Mismatch: ops {actual_max} / watchlist {record['Max Severity']}",
        "Pass" if age_ok else f"Mismatch: ops {actual_avg:.1f} / watchlist {record['Avg Issue Age']}",
        statuses + " | " + dependencies,
    )


wb = Workbook()
wb.remove(wb.active)
wb.calculation.fullCalcOnLoad = True
wb.calculation.forceFullCalc = True

# 1. Workability decision board
ws = wb.create_sheet("Workability Board")
add_title(ws, "Northstar Sprint Workability Board",
          "The control decision comes before risk priority: approved desk work, customer release, controlled exceptions, and holds are separate lanes.", 8)

board_headers = ["Control Lane", "Current Count", "Account IDs", "Can the Desk Work It Now?", "Can the Rep Contact the Customer Now?", "Release Rule", "Owner", "Control Note"]
for c, value in enumerate(board_headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(board_headers), TEAL)

board_rows = [
    ("WORK NOW — CLEAR / PREPARE", 12,
     "A-1001, A-1002, A-1003, A-1004, A-1005, A-1006, A-1007, A-1008, A-1009, A-1010, A-1011, A-1012",
     "Yes — supervisor and owner rep may complete the named evidence and prepare the approved action.",
     "No accounts are customer-ready in the supplied evidence snapshot.",
     "Customer contact only after all four row-level source controls calculate a Cleared release.", SUPERVISOR,
     "These accounts are approved for local execution; rank applies only inside this lane."),
    ("CUSTOMER-READY NOW", '=COUNTIF(\'Approved Work Queue\'!V5:V16,"Cleared")', "See Cleared Contact Path",
     "Only accounts displayed on Cleared Contact Path are released for dated contact.",
     "Only after Release Status calculates Cleared",
     "Current case exports show an open dependency, missing confirmation, or required deliverable for every approved row at package creation.",
     SUPERVISOR, "Reassess from current owner evidence at sprint launch; do not infer clearance from an expected date."),
    ("CONTROLLED EXCEPTION", 4, "A-1013, A-1014, A-1015, A-1017",
     "Internal escalation only.", "No local commitment",
     "Named manager, legal, contact-route, or operational owner must approve a route.", SUPERVISOR,
     "These accounts are physically separated from the work-now queue and carry no execution rank."),
    ("EXCLUDED / HOLD", 2, "A-1016, A-1018",
     "Recordkeeping or future scheduling only.", "No",
     "A-1016 remains closed; A-1018 remains held for the September agency review.", SUPERVISOR,
     "No July recovery outreach."),
]
for r, values in enumerate(board_rows, start=5):
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
style_body(ws, 5, 8, len(board_headers))
for r, color in ((5, LIGHT_BLUE), (6, LIGHT_AMBER), (7, LIGHT_RED), (8, LIGHT_GRAY)):
    for c in range(1, len(board_headers) + 1):
        ws.cell(r, c).fill = PatternFill("solid", fgColor=color)
    ws.cell(r, 1).font = Font(bold=True, color=NAVY)

ws["A11"] = "How an account becomes customer-ready"
ws["A11"].font = Font(size=13, bold=True, color=WHITE)
ws["A11"].fill = PatternFill("solid", fgColor=BLUE)
ws.merge_cells("A11:H11")
release_steps = [
    "1. The July 15 email must name the account for local execution and support the package lane.",
    "2. The directory must show an approved contact route, and the authority matrix must support the exact offer, tier, and value.",
    "3. The operational or billing owner supplies the exact evidence named in Required Evidence Gate; Dana marks only Case Evidence Support.",
    "4. Release Status calculates Cleared only when all four source controls are Supported; only then may the rep make customer contact.",
    "5. A failed control moves to Controlled Exceptions with a named owner and decision date—it is never pushed down the ranking.",
]
for r, value in enumerate(release_steps, start=12):
    ws.cell(r, 1, value)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
    ws.cell(r, 1).alignment = Alignment(wrap_text=True, vertical="top")
    ws.cell(r, 1).border = GRID
    ws.row_dimensions[r].height = 28
set_widths(ws, {"A": 27, "B": 14, "C": 55, "D": 48, "E": 48, "F": 55, "G": 18, "H": 52})
ws.freeze_panes = "A5"
ws.sheet_view.showGridLines = False

# 1. Supervisor control page
ws = wb.create_sheet("Sprint Control")
add_title(ws, "Northstar July Recovery Sprint — Supervisor Work Package",
          "One controlled package for local execution, exceptions, holds, daily checkpoints, and closeout.", 8)
control_rows = [
    ("Supervisor owner", SUPERVISOR, "Approval source", "Elena Whitaker — July 15 approval"),
    ("Execution window", "2026-07-17 through 2026-07-24", "Daily checkpoint", CHECKPOINT),
    ("Approved local execution", 12, "Controlled exceptions", 4),
    ("Excluded / hold", 2, "Operating rule", "A planned date never overrides an uncleared evidence gate."),
]
for r, values in enumerate(control_rows, start=4):
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
    ws.cell(r, 1).font = ws.cell(r, 3).font = Font(bold=True, color=NAVY)
    ws.cell(r, 2).fill = ws.cell(r, 4).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    for c in range(1, 5):
        ws.cell(r, c).border = GRID
        ws.cell(r, c).alignment = Alignment(vertical="top", wrap_text=True)

ws["A10"] = "Supervisor operating sequence"
ws["A10"].font = Font(size=13, bold=True, color=WHITE)
ws["A10"].fill = PatternFill("solid", fgColor=TEAL)
ws.merge_cells("A10:H10")
steps = [
    (1, "Gate", "Confirm the required case evidence is attached; update only Case Evidence Support."),
    (2, "Authorize", "Confirm offer code, tier, requested value, contact route, and customer language."),
    (3, "Assign", "Release the row to the owner rep only after the first two checks pass."),
    (4, "Checkpoint", "Update outreach status, note, next step, and next-step due date by 16:00 Eastern."),
    (5, "Escalate", "Move any failed gate to the exception lane; do not improvise a customer promise."),
    (6, "Close", "On July 24 reconcile model decision, outreach result, open exceptions, and forecast-risk change."),
]
for r, values in enumerate(steps, start=11):
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
    ws.cell(r, 1).font = Font(bold=True, color=WHITE)
    ws.cell(r, 1).fill = PatternFill("solid", fgColor=BLUE)
    ws.cell(r, 2).font = Font(bold=True, color=NAVY)
    for c in range(1, 4):
        ws.cell(r, c).border = GRID
        ws.cell(r, c).alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=8)

ws["A19"] = "Scope and authority guardrails"
ws["A19"].font = Font(size=13, bold=True, color=WHITE)
ws["A19"].fill = PatternFill("solid", fgColor=AMBER)
ws.merge_cells("A19:H19")
guardrails = [
    "Use only MG-15, CR-10K, RH-60, SW-PKG, or RPT-COMP where the account row passes its authority check.",
    "Do not promise future performance, unverified credits, new product terms, legal terms, or an uncleared operational action.",
    "MGR-REQ is an escalation route, not a customer offer.",
    "Cobalt Cloud (A-1003): local authority covers the verified CR-10K credit only; RH-60 is not eligible for Key Growth.",
]
for r, rule in enumerate(guardrails, start=20):
    ws.cell(r, 1, "• " + rule)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
    ws.cell(r, 1).alignment = Alignment(wrap_text=True, vertical="top")
    ws.cell(r, 1).border = GRID
    ws.row_dimensions[r].height = 30
set_widths(ws, {"A": 18, "B": 27, "C": 18, "D": 42, "E": 13, "F": 13, "G": 13, "H": 13})
ws.sheet_view.showGridLines = False

# 2. Approved work queue
ws = wb.create_sheet("Approved Work Queue")
headers = [
    "Desk Work Lane", "Execution Rank", "Conditional Contact Target", "Account ID", "Account Name", "Owner Rep", "Supervisor",
    "Renewal Date", "Days to Renewal", "Renewal Exposure", "Forecast Variance", "Priority Score",
    "Primary Issue", "Case IDs", "Operational Owner(s)", "Expected Clearance", "Required Evidence Gate",
    "Approval Email Support", "Contact Route Support", "Case Evidence Support", "Offer Authority Support",
    "Release Status", "Released Contact Date", "Approved Offer", "Authority Limit", "Requested Value", "Authority Check",
    "Approved Contact", "Contact Role", "Preferred Channel", "Timezone", "Contact Restriction",
    "Supervisor-Approved Customer Action", "Customer Language Guardrail", "Current Rep Commitment",
    "Commitment Due", "Outreach Status", "Checkpoint Note", "Next Step", "Next Step Due"
]
add_title(ws, "Approved Local Execution Queue",
          "Release is formula-controlled: approval email + contact route + case evidence + offer authority must all support the exact action.", len(headers))
for c, value in enumerate(headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers))

approved = [r for r in watchlist if r["Approval Status"] == "Approved execution"]
approved.sort(key=lambda r: execution_plan[r["Account ID"]][0])
for row_num, record in enumerate(approved, start=5):
    aid = record["Account ID"]
    rank, planned, gate, language = execution_plan[aid]
    contact = primary_contact(aid)
    case_ids, owners, expected = case_summary(aid)
    commitment, commitment_due = current_commitment(aid)
    action = record["Decision Action"]
    authority_check = "Within local authority after gate clears"
    if aid == "A-1003":
        action = "Confirm verified invoice credit only; escalate any rate-hold request"
        authority_check = "CR-10K only; RH-60 is not eligible for Key Growth"
    approval_support = "Supported — July 15 approval"
    contact_support = f'Supported — {contact["Contact Name"]} / Approved'
    case_support = "Pending — evidence not yet confirmed"
    offer_support = f'Supported — {record["Approved Offer"]} / {record["Account Tier"]} / ${int(record["Requested Recovery Value"] or 0):,}'
    values = [
        "WORK NOW — CLEAR / PREPARE", rank, planned, aid, record["Account Name"], record["Owner Rep"], SUPERVISOR,
        as_date(record["Renewal Date"]), None, record["Renewal Exposure"], record["Forecast Variance"],
        record["Priority Score"], record["Primary Issue"], case_ids, owners,
        datetime.strptime(expected, "%Y-%m-%d").date() if expected else None, gate,
        approval_support, contact_support, case_support, offer_support, None, None,
        record["Approved Offer"], offer_caps[record["Approved Offer"]], record["Requested Recovery Value"],
        authority_check, contact["Contact Name"], contact["Contact Role"], contact["Preferred Channel"],
        contact["Timezone"], f'{contact["Outreach Permission"]}; {contact["Sensitivity Flag"]}; {contact["Unavailable Window"]}',
        action, language, commitment,
        datetime.strptime(commitment_due, "%Y-%m-%d").date(), "Not started", "", action, planned
    ]
    for c, value in enumerate(values, start=1):
        ws.cell(row_num, c, value)
    ws.cell(row_num, 9, f"=H{row_num}-C{row_num}")
    ws.cell(row_num, 22, f'=IF(COUNTIF(R{row_num}:U{row_num},"Supported*")=4,"Cleared","Blocked")')
    ws.cell(row_num, 23, f'=IF(V{row_num}="Cleared",C{row_num},"")')

end_row = 4 + len(approved)
style_body(ws, 5, end_row, len(headers))
for row in range(5, end_row + 1):
    for col in (3, 8, 16, 23, 36, 40):
        ws.cell(row, col).number_format = "mmm d"
    for col in (10, 11, 26):
        ws.cell(row, col).number_format = '$#,##0;[Red]-$#,##0'
    ws.cell(row, 12).number_format = "0.0"
    ws.cell(row, 1).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    ws.cell(row, 20).fill = PatternFill("solid", fgColor=LIGHT_AMBER)
    ws.cell(row, 22).fill = PatternFill("solid", fgColor=LIGHT_RED)
    ws.cell(row, 23).fill = PatternFill("solid", fgColor=LIGHT_GRAY)
    ws.cell(row, 27).fill = PatternFill("solid", fgColor=LIGHT_GREEN if row != 9 else LIGHT_AMBER)
    ws.cell(row, 37).fill = PatternFill("solid", fgColor=LIGHT_BLUE)

case_dv = DataValidation(type="list", formula1='"Pending — evidence not yet confirmed,Supported — evidence attached,Not supported — escalate"', allow_blank=False)
outreach_dv = DataValidation(type="list", formula1='"Not started,Prepared,Sent,Meeting booked,Completed,Hold"', allow_blank=False)
ws.add_data_validation(case_dv)
ws.add_data_validation(outreach_dv)
case_dv.add(f"T5:T{end_row}")
outreach_dv.add(f"AK5:AK{end_row}")
ws.conditional_formatting.add(f"V5:V{end_row}", FormulaRule(formula=["V5=\"Cleared\""], fill=PatternFill("solid", fgColor=LIGHT_GREEN)))
ws.conditional_formatting.add(f"V5:V{end_row}", FormulaRule(formula=["V5=\"Blocked\""], fill=PatternFill("solid", fgColor=LIGHT_RED)))
ws.conditional_formatting.add(f"W5:W{end_row}", FormulaRule(formula=['AND(W5<>"",V5<>"Cleared")'], fill=PatternFill("solid", fgColor=LIGHT_RED)))
ws.conditional_formatting.add(f"AK5:AK{end_row}", FormulaRule(formula=["AK5=\"Completed\""], fill=PatternFill("solid", fgColor=LIGHT_GREEN)))
add_table(ws, f"A4:AN{end_row}", "ApprovedWorkQueue", "TableStyleMedium2")
ws.freeze_panes = "H5"
ws.auto_filter.ref = f"A4:AN{end_row}"
set_widths(ws, {
    "A": 27, "B": 11, "C": 13, "D": 11, "E": 27, "F": 16, "G": 16, "H": 12, "I": 11,
    "J": 16, "K": 15, "L": 12, "M": 22, "N": 20, "O": 30, "P": 14, "Q": 48,
    "R": 29, "S": 34, "T": 34, "U": 38, "V": 16, "W": 18, "X": 14, "Y": 32,
    "Z": 15, "AA": 40, "AB": 24, "AC": 23, "AD": 18, "AE": 12, "AF": 43,
    "AG": 45, "AH": 55, "AI": 38, "AJ": 13, "AK": 16, "AL": 38, "AM": 42, "AN": 13
})
ws.sheet_view.showGridLines = False

# 3. Cleared-only dated contact path
ws = wb.create_sheet("Cleared Contact Path")
headers = ["Released Contact Date", "Account ID", "Account Name", "Owner Rep", "Approved Contact",
           "Preferred Channel", "Supervisor-Approved Customer Action", "Release Status", "Outreach Status"]
add_title(ws, "Cleared-Only Dated Contact Path",
          "This view populates only when all four source controls calculate Release Status = Cleared. Blocked, controlled, and held accounts cannot appear here.", len(headers))
for c, value in enumerate(headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), GREEN)

for r in range(5, 17):
    source_columns = {1: "W", 2: "D", 3: "E", 4: "F", 5: "AB", 6: "AD", 7: "AG", 8: "V", 9: "AK"}
    for col, source_col in source_columns.items():
        ws.cell(r, col, f'=IF(\'Approved Work Queue\'!$V${r}="Cleared",\'Approved Work Queue\'!${source_col}${r},"")')
    ws.cell(r, 1).number_format = "ddd, mmm d"
style_body(ws, 5, 16, len(headers))
for r in range(5, 17):
    ws.row_dimensions[r].height = 32
ws.freeze_panes = "A5"
set_widths(ws, {"A": 20, "B": 12, "C": 29, "D": 17, "E": 25, "F": 20, "G": 52, "H": 16, "I": 16})
ws.sheet_view.showGridLines = False

# 4. Daily runbook
ws = wb.create_sheet("Daily Runbook")
headers = ["Date", "Supervisor Focus", "Pre-contact Clearance Work", "Cleared Customer Contacts Only", "Exception / Hold Check", "16:00 ET Checkpoint Deliverable", "Completion Status", "Supervisor Note"]
add_title(ws, "July 17–24 Daily Runbook", "Column C prepares accounts against their conditional contact targets. Only accounts appearing on Cleared Contact Path may enter column D.", len(headers))
for c, value in enumerate(headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), TEAL)
runbook = [
    (date(2026, 7, 17), "Launch: validate evidence, authorize only cleared rows, assign rep actions.",
     "A-1001, A-1003, A-1004, A-1009", "None until rows appear on Cleared Contact Path", "A-1014 manager credit decision",
     "Gate status, released actions, blocked items, and owner/due date."),
    (date(2026, 7, 20), "Clear weekend operational dependencies and run highest-risk outreach.",
     "A-1005, A-1012, A-1007, A-1006", "Use Cleared Contact Path only", "A-1013 contact-route decision",
     "Outreach result, customer response, missed gates, and updated forecast risk."),
    (date(2026, 7, 21), "Complete reporting-package actions and evidence-only follow-up.",
     "A-1002, A-1011, A-1008", "Use Cleared Contact Path only", "A-1015 product-feed status",
     "Reports delivered, follow-ups booked, exception changes, and forecast risk."),
    (date(2026, 7, 22), "Complete rate-hold action and rescue any slipped approved account.",
     "A-1010 plus slipped approved rows", "Use Cleared Contact Path only", "A-1015 feed-fix decision",
     "Completed actions, rescue assignments, outstanding customer decisions."),
    (date(2026, 7, 23), "Resolve follow-ups and prepare closeout reconciliation.",
     "All open approved rows", "Use Cleared Contact Path only", "A-1017 legal route status",
     "Open actions with owners/dates; draft decision and forecast reconciliation."),
    (date(2026, 7, 24), "Final follow-up and sprint closeout.",
     "A-1001 renewal checkpoint; all remaining approved rows", "Use Cleared Contact Path only", "All four exceptions; confirm A-1016/A-1018 remain held",
     "Final model decision, outreach status, exceptions, and forecast-risk change."),
]
for r, values in enumerate(runbook, start=5):
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
    ws.cell(r, 7, "Not started")
    ws.cell(r, 1).number_format = "ddd, mmm d"
style_body(ws, 5, 10, len(headers))
status_dv = DataValidation(type="list", formula1='"Not started,In progress,Complete,Blocked"', allow_blank=False)
ws.add_data_validation(status_dv)
status_dv.add("G5:G10")
for r in range(5, 11):
    ws.cell(r, 7).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
add_table(ws, "A4:H10", "DailyRunbook", "TableStyleMedium4")
ws.freeze_panes = "A5"
set_widths(ws, {"A": 15, "B": 42, "C": 40, "D": 36, "E": 42, "F": 52, "G": 16, "H": 48})
ws.sheet_view.showGridLines = False

# 4. Controlled exceptions
ws = wb.create_sheet("Controlled Exceptions")
headers = ["Control Lane", "Account ID", "Account Name", "Owner Rep", "Supervisor", "Renewal Date", "Renewal Exposure",
           "Forecast Variance", "Primary Issue", "Exception Reason", "Escalation Owner", "Decision Due", "Required Route",
           "Customer Contact Rule", "Exception Status", "Checkpoint Note", "Next Step"]
add_title(ws, "Controlled Exception Queue",
          "Visible to the supervisor, but not releasable as local customer commitments.", len(headers))
for c, value in enumerate(headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), AMBER)
exceptions = [r for r in watchlist if r["Approval Status"] == "Controlled exception"]
exceptions.sort(key=lambda r: r["Rank"])
for r, record in enumerate(exceptions, start=5):
    aid = record["Account ID"]
    owner, due, route = exception_owner[aid]
    values = ["CONTROLLED — NO LOCAL COMMITMENT", aid, record["Account Name"], record["Owner Rep"], SUPERVISOR,
              as_date(record["Renewal Date"]), record["Renewal Exposure"], record["Forecast Variance"],
              record["Primary Issue"], record["Exception Reason"], owner, due, route,
              "No local commitment; evidence-only contact where directory permits.", "Open", "", route]
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
    for c in (6, 12): ws.cell(r, c).number_format = "mmm d"
    for c in (7, 8): ws.cell(r, c).number_format = '$#,##0;[Red]-$#,##0'
style_body(ws, 5, 8, len(headers))
exception_dv = DataValidation(type="list", formula1='"Open,Pending owner,Approved route,Remain on hold,Closed"', allow_blank=False)
ws.add_data_validation(exception_dv)
exception_dv.add("O5:O8")
for r in range(5, 9): ws.cell(r, 15).fill = PatternFill("solid", fgColor=LIGHT_AMBER)
add_table(ws, "A4:Q8", "ExceptionQueue", "TableStyleMedium12")
ws.freeze_panes = "F5"
set_widths(ws, {"A": 10, "B": 11, "C": 28, "D": 16, "E": 16, "F": 13, "G": 16, "H": 15,
                "I": 24, "J": 34, "K": 28, "L": 13, "M": 52, "N": 45, "O": 17, "P": 42, "Q": 48})
ws.sheet_view.showGridLines = False

# 5. Explicit holds
ws = wb.create_sheet("Excluded and Holds")
headers = ["Account ID", "Account Name", "Owner Rep", "Status", "Reason", "Permitted Supervisor Action", "Review Date", "Note"]
add_title(ws, "Excluded and Hold Records", "No sprint outreach or local commercial action.", len(headers))
for c, value in enumerate(headers, start=1): ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), GRAY)
hold_rows = [
    ("A-1016", watch_by_id["A-1016"]["Account Name"], watch_by_id["A-1016"]["Owner Rep"], "Excluded — closed renewal",
     "Issue resolved and renewal signed.", "Retain closure evidence only.", date(2026, 7, 24), "No outreach."),
    ("A-1018", watch_by_id["A-1018"]["Account Name"], watch_by_id["A-1018"]["Owner Rep"], "Customer hold",
     "Agency requested September consolidation review.", "Preserve hold and schedule September rebrief only.", date(2026, 9, 2), "No July outreach."),
]
for r, values in enumerate(hold_rows, start=5):
    for c, value in enumerate(values, start=1): ws.cell(r, c, value)
    ws.cell(r, 7).number_format = "mmm d"
style_body(ws, 5, 6, len(headers))
add_table(ws, "A4:H6", "HoldQueue", "TableStyleMedium15")
set_widths(ws, {"A": 11, "B": 30, "C": 16, "D": 24, "E": 42, "F": 48, "G": 14, "H": 30})
ws.sheet_view.showGridLines = False

# 6. Per-account source reconciliation
ws = wb.create_sheet("Source Reconciliation")
headers = [
    "Account ID", "Account Name", "July 15 Approval Scope", "Package Control Lane", "Scope Match",
    "CRM Lifecycle State", "CRM Opportunity Approval", "Owner Match", "Renewal Date Match",
    "Exposure Match", "Requested Value Match", "Watchlist Risk Rank", "Watchlist Offer",
    "Offer / Authority Check", "Ops Case Count", "Ops Max Severity", "Ops Average Age",
    "Ops Status and Dependencies", "Contact Route", "Contact Control Check",
    "Conditional Contact Target", "Initial Release State", "Overall Reconciliation",
    "Required Supervisor Disposition", "Correction / Control Note"
]
add_title(ws, "Account-by-Account Source Reconciliation",
          "Reconciles the package to CRM health, opportunity exposure, ops cases, contact authority, offer authority, the watchlist, and the July 15 approval.", len(headers))
for c, value in enumerate(headers, start=1):
    ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), NAVY)

for r, record in enumerate(sorted(watchlist, key=lambda item: item["Account ID"]), start=5):
    aid = record["Account ID"]
    crm = accounts[aid]
    opp = opportunities[aid]
    approval_scope = july_scope(aid)
    lane = package_lane(aid)
    expected_watch_status = (
        "Approved execution" if aid in approved_scope_ids else
        "Controlled exception" if aid in controlled_scope_ids else
        "Excluded - no action"
    )
    lifecycle_ok = (
        crm["Lifecycle State"] == "Approved execution" if aid in approved_scope_ids else
        crm["Lifecycle State"] == "Controlled exception" if aid in controlled_scope_ids else
        str(crm["Lifecycle State"]).startswith("Excluded")
    )
    opportunity_ok = (
        opp["Approval State"] == "Approved execution" if aid in approved_scope_ids else
        opp["Approval State"] == "Controlled exception" if aid in controlled_scope_ids else
        opp["Approval State"] == "Excluded - no action"
    )
    scope_ok = record["Approval Status"] == expected_watch_status and lifecycle_ok and opportunity_ok
    owner_ok = record["Owner Rep"] == crm["Owner Rep"] == opp["Owner Rep"]
    renewal_ok = as_date(record["Renewal Date"]) == as_date(crm["Renewal Date"]) == as_date(opp["Renewal Date"])
    exposure_ok = float(record["Renewal Exposure"] or 0) == float(opp["Renewal Exposure"] or 0)
    requested_ok = float(record["Requested Recovery Value"] or 0) == float(crm["Requested Recovery Value"] or 0)
    ops_count, ops_severity, ops_age, ops_detail = ops_reconciliation(record)
    contact_route, contact_check = contact_route_summary(aid)
    offer_check = offer_reconciliation(record)

    if aid in approved_scope_ids:
        target = execution_plan[aid][1]
        release_state = "Pending evidence — no customer contact"
        disposition = execution_plan[aid][2]
        overall = "RECONCILED — PENDING EVIDENCE"
    elif aid in controlled_scope_ids:
        target = None
        release_state = "Controlled — no local commitment"
        disposition = exception_owner[aid][2]
        overall = "RECONCILED — CONTROLLED"
    else:
        target = None
        release_state = "Excluded / hold — no July contact"
        disposition = record["Decision Action"]
        overall = "RECONCILED — EXCLUDED / HOLD"

    checks = [scope_ok, owner_ok, renewal_ok, exposure_ok, requested_ok,
              ops_count == "Pass", ops_severity == "Pass", ops_age == "Pass",
              not offer_check.startswith("Mismatch"), contact_check != "Mismatch"]
    correction_note = ""
    if aid == "A-1003":
        correction_note = "Watchlist action mixed CR-10K with an ineligible RH-60 rate hold. Package restricts local action to CR-10K and treats CASE-7204 as the controlling evidence date."
    elif aid == "A-1013":
        correction_note = "Offer fit alone does not clear the account; final-approver permission remains pending."
    elif aid == "A-1014":
        correction_note = "The $42,000 request exceeds local authority and remains MGR-REQ."
    elif aid == "A-1015":
        correction_note = "SW-PKG fits tier/value, but the product-feed blocker prevents local release."
    elif aid == "A-1017":
        correction_note = "Enterprise data-use terms remain with manager/legal; executive contact is prohibited."
    elif aid == "A-1018":
        correction_note = "Customer-requested September hold overrides risk priority."
    if not all(checks):
        overall = "RECONCILIATION MISMATCH — REVIEW"
        mismatch_labels = [
            label for label, ok in zip(
                ["scope", "owner", "renewal", "exposure", "requested value", "ops count", "ops severity", "ops age", "offer", "contact"],
                checks
            ) if not ok
        ]
        correction_note = (correction_note + " | " if correction_note else "") + "Mismatch: " + ", ".join(mismatch_labels)

    values = [
        aid, record["Account Name"], approval_scope, lane, "Pass" if scope_ok else "Mismatch",
        crm["Lifecycle State"], opp["Approval State"], "Pass" if owner_ok else "Mismatch",
        "Pass" if renewal_ok else "Mismatch", "Pass" if exposure_ok else "Mismatch",
        "Pass" if requested_ok else "Mismatch", record["Rank"], record["Approved Offer"] or "None",
        offer_check, ops_count, ops_severity, ops_age, ops_detail, contact_route, contact_check,
        target, release_state, overall, disposition, correction_note
    ]
    for c, value in enumerate(values, start=1):
        ws.cell(r, c, value)
    ws.cell(r, 21).number_format = "mmm d"

recon_end = 4 + len(watchlist)
style_body(ws, 5, recon_end, len(headers))
for r in range(5, recon_end + 1):
    status = ws.cell(r, 23).value
    color = LIGHT_GREEN if "PENDING EVIDENCE" in status else LIGHT_AMBER if "CONTROLLED" in status else LIGHT_GRAY
    if "MISMATCH" in status:
        color = LIGHT_RED
    ws.cell(r, 23).fill = PatternFill("solid", fgColor=color)
    ws.cell(r, 23).font = Font(bold=True, color=NAVY)
add_table(ws, f"A4:Y{recon_end}", "SourceReconciliation", "TableStyleMedium2")
ws.freeze_panes = "F5"
set_widths(ws, {
    "A": 11, "B": 29, "C": 24, "D": 37, "E": 12, "F": 26, "G": 25, "H": 12,
    "I": 15, "J": 14, "K": 18, "L": 13, "M": 15, "N": 42, "O": 16, "P": 16,
    "Q": 16, "R": 65, "S": 34, "T": 45, "U": 20, "V": 34, "W": 34, "X": 58, "Y": 68
})
ws.sheet_view.showGridLines = False

# 7. Authority quick reference
ws = wb.create_sheet("Authority Reference")
headers = ["Offer Code", "Issue Basis", "Eligible Tiers", "Authority Limit", "Approval", "Permitted Use", "If Evidence Is Missing"]
add_title(ws, "Local Recovery Authority — Quick Reference",
          "Effective July 15–September 30, 2026. The source authority document remains controlling.", len(headers))
for c, value in enumerate(headers, start=1): ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), GREEN)
missing_evidence_route = {
    "MG-15": "Hold", "CR-10K": "Hold", "RH-60": "Exception",
    "SW-PKG": "Hold", "RPT-COMP": "Hold", "MGR-REQ": "Manager / functional owner"
}
authority_rows = [
    (code, source["Issue Basis"], source["Eligible Tiers"], source["Max Value"],
     source["Approval"], source["Permitted Use"], missing_evidence_route[code])
    for code, source in authority_source.items()
]
for r, values in enumerate(authority_rows, start=5):
    for c, value in enumerate(values, start=1): ws.cell(r, c, value)
style_body(ws, 5, 10, len(headers))
add_table(ws, "A4:G10", "AuthorityReference", "TableStyleMedium4")
set_widths(ws, {"A": 14, "B": 48, "C": 33, "D": 25, "E": 18, "F": 58, "G": 25})
ws.sheet_view.showGridLines = False

# 7. Source control and known correction
ws = wb.create_sheet("Source Control")
headers = ["Source", "As-of / Effective Date", "Use in Package", "Control Note"]
add_title(ws, "Source Lineage and Package Corrections", "Source dates are preserved so the supervisor can identify stale evidence before release.", len(headers))
for c, value in enumerate(headers, start=1): ws.cell(4, c, value)
style_header(ws, 4, 1, len(headers), NAVY)
source_rows = [
    (str(watchlist_path.relative_to(ROOT)), "Review 2026-07-14", "Risk scores, tiers, proposed actions", "Global risk rank replaced by an execution-only rank plus separate exception/hold lanes."),
    (str(crm_path.relative_to(ROOT)), "Export 2026-07-13", "Account, renewal, approval, forecast context", "Refresh changed renewal or approval state before outreach."),
    (str(cases_path.relative_to(ROOT)), "Queue 2026-07-13; some updates through 2026-07-14", "Case owner, dependency, expected resolution", "Gate clearance requires current owner evidence, not the file date alone."),
    (str(contacts_path.relative_to(ROOT)), "Directory 2026-07-10", "Approved contact route and restrictions", "Reconfirm any pending or restricted route."),
    (str(activity_path.relative_to(ROOT)), "Through 2026-07-14", "Commitments, due dates, next steps", "Checkpoint fields in this package become the sprint record."),
    ("data/commercial_options/local_recovery_offer_authority_q3_2026.docx", "Effective 2026-07-15 to 2026-09-30", "Offer codes, caps, evidence gates, escalation rules", "Controlling authority source."),
    ("data/approvals/recovery_sprint_scope_approval_2026-07-15.eml", "Approved 2026-07-15", "Supervisor, execution window, account scope", "Controlling sprint approval source."),
    ("Package correction: A-1003", "Applied in this workbook", "Local action and guardrail", "Removed unsupported RH-60 language; Key Growth is not RH-60 eligible. CR-10K only unless separately approved."),
]
for r, values in enumerate(source_rows, start=5):
    for c, value in enumerate(values, start=1): ws.cell(r, c, value)
style_body(ws, 5, 12, len(headers))
add_table(ws, "A4:D12", "SourceControl", "TableStyleMedium2")
set_widths(ws, {"A": 70, "B": 30, "C": 45, "D": 72})
ws.sheet_view.showGridLines = False

# Workbook-wide print and view settings.
for ws in wb.worksheets:
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_setup.orientation = "landscape"
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5
    ws.sheet_view.zoomScale = 85

OUTPUT.mkdir(parents=True, exist_ok=True)
wb.save(OUTPUT_FILE)
print(OUTPUT_FILE)
