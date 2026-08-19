from collections import defaultdict
from datetime import datetime
from email import policy
from email.parser import BytesParser
from pathlib import Path
import csv
import json

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Protection, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.table import Table, TableStyleInfo


BASE = Path("/app/data/care_transition")
OUTPUT = Path("/app/med_rec.xlsx")


# Products that the controlled reference allows to be reconciled in one row are
# grouped together. Explicit non-substitutions (for example, metoprolol
# tartrate/succinate and atorvastatin/rosuvastatin) remain separate rows.
MEDICATIONS = [
    ("apixaban", "Apixaban 5 mg tablet", ["apixaban", "eliquis"]),
    ("warfarin", "Warfarin 5 mg tablet", ["warfarin"]),
    ("metoprolol_tartrate", "Metoprolol tartrate 50 mg tablet", ["metoprolol tartrate", "lopressor"]),
    ("metoprolol_succinate", "Metoprolol succinate ER 100 mg tablet", ["metoprolol succinate"]),
    ("lisinopril", "Lisinopril 20 mg tablet", ["lisinopril"]),
    ("losartan", "Losartan 50 mg tablet", ["losartan"]),
    (
        "basal_insulin",
        "Insulin glargine U-100 pen",
        ["lantus", "basaglar", "semglee", "insulin glargine-yfgn", "insulin glargine u-100"],
    ),
    ("insulin_lispro", "Insulin lispro U-100 pen", ["insulin lispro", "humalog"]),
    ("metformin", "Metformin ER 500 mg tablet", ["metformin"]),
    ("furosemide", "Furosemide tablet", ["furosemide", "water pill"]),
    ("potassium", "Potassium chloride ER 20 mEq tablet", ["potassium chloride", "klor-con"]),
    ("spironolactone", "Spironolactone 25 mg tablet", ["spironolactone"]),
    ("aspirin", "Aspirin 81 mg EC tablet", ["aspirin", "baby aspirin"]),
    ("clopidogrel", "Clopidogrel 75 mg tablet", ["clopidogrel"]),
    ("albuterol", "Albuterol HFA 90 mcg inhaler", ["albuterol", "ventolin", "proair"]),
    (
        "fluticasone_salmeterol",
        "Fluticasone-salmeterol 250/50 Diskus",
        ["fluticasone-salmeterol", "advair"],
    ),
    ("prednisone", "Prednisone 10 mg taper", ["prednisone"]),
    ("amox_clav", "Amoxicillin-clavulanate 875/125 mg tablet", ["amoxicillin-clavulanate"]),
    ("naproxen", "Naproxen 500 mg tablet", ["naproxen"]),
    ("acetaminophen", "Acetaminophen 500 mg tablet", ["acetaminophen", "tylenol"]),
    ("ppi", "PPI: pantoprazole 40 mg / omeprazole 20 mg", ["pantoprazole", "omeprazole"]),
    ("atorvastatin", "Atorvastatin 40 mg tablet", ["atorvastatin"]),
    ("rosuvastatin", "Rosuvastatin 20 mg tablet", ["rosuvastatin"]),
    ("levothyroxine", "Levothyroxine 75 mcg tablet", ["levothyroxine"]),
    ("sertraline", "Sertraline 50 mg tablet", ["sertraline"]),
    ("trazodone", "Trazodone 50 mg tablet", ["trazodone"]),
    ("vitamin_d", "Vitamin D3 2,000 unit capsule", ["vitamin d"]),
    ("fish_oil", "Fish oil 1,000 mg capsule", ["fish oil"]),
    ("multivitamin", "Multivitamin tablet", ["multivitamin"]),
    ("docusate", "Docusate 100 mg capsule", ["docusate"]),
    ("senna", "Senna 8.6 mg tablet", ["senna"]),
    ("oxycodone", "Oxycodone 5 mg tablet", ["oxycodone"]),
    ("gabapentin", "Gabapentin 300 mg capsule", ["gabapentin"]),
    ("nitroglycerin", "Nitroglycerin 0.4 mg SL tablet", ["nitroglycerin", "nitro tabs"]),
    ("magnesium", "Magnesium oxide 400 mg tablet", ["magnesium oxide"]),
]

DISPLAY = {key: display for key, display, _ in MEDICATIONS}
ALIASES = {key: aliases for key, _, aliases in MEDICATIONS}


def med_key(text):
    value = (text or "").lower().replace("‑", "-")
    for key, _, aliases in MEDICATIONS:
        if any(alias in value for alias in aliases):
            return key
    return None


def compact(parts):
    return " | ".join(str(part).strip() for part in parts if part not in (None, ""))


def read_sheet_rows(filename, sheet_name):
    wb = load_workbook(BASE / filename, read_only=True, data_only=True)
    ws = wb[sheet_name]
    rows = list(ws.iter_rows(values_only=True))
    headers = [str(x) for x in rows[0]]
    return [dict(zip(headers, row)) for row in rows[1:]]


admission = read_sheet_rows("admission_history_snapshot.xlsx", "Medication_History")
portal = read_sheet_rows("patient_list_snapshot.xlsx", "Portal_Medications")
profile = [
    row
    for row in read_sheet_rows("pharmacy_profile_snapshot.xlsx", "Profile_Medications")
    if row["case_id"] == "CT-2026-0713-18"
]
dispenses = [
    row
    for row in read_sheet_rows("pharmacy_profile_snapshot.xlsx", "Dispense_History")
    if row["case_id"] == "CT-2026-0713-18"
]

with (BASE / "reconciliation_status_extract.csv").open(
    newline="", encoding="utf-8-sig"
) as handle:
    queue = list(csv.DictReader(handle))

with (BASE / "allergy_status_extract.csv").open(
    newline="", encoding="utf-8-sig"
) as handle:
    allergies = list(csv.DictReader(handle))

manifest = json.loads((BASE / "transition_packet_manifest.json").read_text())
source_exported_at = {
    source["relative_path"]: source["exported_at"]
    for source in manifest["source_register"]
}

summary_doc = Document(BASE / "transition_summary.docx")
summary_rows = []
for table in summary_doc.tables[1:3]:
    headers = [cell.text for cell in table.rows[0].cells]
    for row in table.rows[1:]:
        summary_rows.append(dict(zip(headers, [cell.text for cell in row.cells])))

email_msg = BytesParser(policy=policy.default).parse(
    (BASE / "clarification_thread_july.eml").open("rb")
)
email_text = next(
    part.get_content()
    for part in email_msg.walk()
    if part.get_content_type() == "text/plain"
)


by_admission = defaultdict(list)
for row in admission:
    key = med_key(row["medication_label"])
    if key:
        by_admission[key].append(
            compact(
                [
                    row["source_row_id"],
                    row["list_version"],
                    row["medication_label"],
                    row["strength"],
                    row["sig"],
                    row["reported_status"],
                    f"captured {row['captured_at']}",
                ]
            )
        )

by_portal = defaultdict(list)
for row in portal:
    key = med_key(row["medication_label"])
    if key:
        by_portal[key].append(
            compact(
                [
                    row["source_row_id"],
                    row["medication_label"],
                    row["dose_text"],
                    row["schedule_text"],
                    row["patient_status"],
                    row["patient_comment"],
                ]
            )
        )

by_profile = defaultdict(list)
profile_row_to_key = {}
for row in profile:
    key = med_key(row["medication_label"])
    if key:
        profile_row_to_key[row["source_row_id"]] = key
        by_profile[key].append(
            compact(
                [
                    row["source_row_id"],
                    row["medication_label"],
                    f"profile={row['profile_status']}",
                    f"fill={row['fill_status']}",
                    f"last_fill={row['last_fill_date']}",
                    row["sig"],
                    row["profile_note"],
                ]
            )
        )

by_dispense = defaultdict(list)
for row in dispenses:
    key = profile_row_to_key.get(row["profile_row_id"]) or med_key(row["medication_label"])
    if key:
        by_dispense[key].append(row)

by_summary = defaultdict(list)
summary_disposition = {}
summary_directions = {}
for row in summary_rows:
    key = med_key(row["Medication"])
    if key:
        summary_disposition[key] = {
            "CONTINUE": "KEEP",
            "CHANGE": "CHANGE",
            "STOP": "STOP",
        }[row["Disposition"].upper()]
        summary_directions[key] = compact([row["Medication"], row["Dose / directions"]])
        by_summary[key].append(
            compact(
                [
                    row["Line ID"],
                    row["Disposition"],
                    row["Medication"],
                    row["Dose / directions"],
                    row["Medication-note text"],
                ]
            )
        )

by_queue = defaultdict(list)
for row in queue:
    key = med_key(row["medication_or_scope"])
    if key:
        by_queue[key].append(
            compact(
                [
                    row["queue_id"],
                    row["event_time"],
                    row["proposed_action"],
                    f"{row['queue_state']}/{row['revision_status']}",
                    row["note"],
                ]
            )
        )


clarifications = {
    "metoprolol_tartrate": "2026-07-14 15:10 — Do not carry tartrate forward.",
    "metoprolol_succinate": "2026-07-14 15:10 — Use succinate ER 100 mg once daily from signed summary.",
    "ppi": "2026-07-14 15:10 — Outpatient formulary substitution approved; no do-not-substitute note.",
    "potassium": "2026-07-14 15:10 — BMP at 14:36 reportedly K 4.2; continue ER 20 mEq once daily.",
    "spironolactone": "2026-07-14 15:10 — Not confirmed; ask cardiology before dispensing.",
    "magnesium": "2026-07-14 15:10 — Needs an answer from discharging service.",
}

allergy_notes = {
    "lisinopril": "ACTIVE HIGH allergy: lisinopril—angioedema; verified 2026-07-12 14:22.",
    "amox_clav": "ACTIVE HIGH penicillin-class allergy: hives; verified 2026-07-10 12:44.",
    "naproxen": "ACTIVE HIGH naproxen allergy: hives; verified 2026-07-10 12:47.",
}

risk = {
    "apixaban": "HIGH",
    "warfarin": "HIGH",
    "basal_insulin": "HIGH",
    "insulin_lispro": "HIGH",
    "lisinopril": "HIGH",
    "amox_clav": "HIGH",
    "naproxen": "HIGH",
    "oxycodone": "HIGH",
    "metoprolol_tartrate": "MED",
    "metoprolol_succinate": "MED",
    "losartan": "MED",
    "furosemide": "MED",
    "potassium": "MED",
    "spironolactone": "MED",
    "nitroglycerin": "MED",
    "magnesium": "MED",
}

proposed = {}
for key, _, _ in MEDICATIONS:
    proposed[key] = summary_disposition.get(key, "CLARIFY")
proposed.update(
    {
        "spironolactone": "CLARIFY",
        "magnesium": "CLARIFY",
        "ppi": "CHANGE",
        "potassium": "KEEP",
    }
)

status = {}
gap = {}
for key, _, _ in MEDICATIONS:
    if key in ("spironolactone", "magnesium"):
        status[key] = "OPEN"
    elif key in summary_disposition or key in clarifications:
        status[key] = "RESOLVED"
    else:
        status[key] = "REVIEW"

gap.update(
    {
        "spironolactone": "Final summary silent; obtain cardiology decision before dispensing.",
        "magnesium": "Final summary silent; obtain discharging-service decision.",
        "lisinopril": "HIGH-RISK mismatch: pharmacy profile remains active/ready_to_fill despite severe allergy and stop order.",
        "clopidogrel": "Pharmacy profile remains active despite explicit discharge stop.",
        "atorvastatin": "Pharmacy profile remains active despite stop/change to rosuvastatin.",
        "trazodone": "Pharmacy profile remains active despite explicit discharge stop.",
        "basal_insulin": "Latest pharmacy row reflects 28 units; signed discharge dose is 30 units.",
        "furosemide": "Latest pharmacy row reflects 20 mg; signed discharge dose is 40 mg.",
        "metoprolol_succinate": "New formulation is absent from latest pharmacy profile.",
        "losartan": "New discharge medication is absent from latest pharmacy profile.",
        "rosuvastatin": "New discharge medication is absent from latest pharmacy profile.",
        "senna": "New discharge bowel regimen is absent from latest pharmacy profile.",
        "potassium": "Resolved by email/queue, but final signed summary was not amended and preliminary queue row remains open.",
        "warfarin": "Stop is final, but an obsolete preliminary queue row remains open.",
        "metoprolol_tartrate": "Stop/change is final, but an obsolete preliminary queue row remains open.",
        "ppi": "Resolved operationally through formulary substitution; no amended signed summary.",
    }
)


wb = Workbook()
ws = wb.active
ws.title = "Med_Reconciliation"

title_fill = PatternFill("solid", fgColor="1F4E78")
header_fill = PatternFill("solid", fgColor="D9EAF7")
subheader_fill = PatternFill("solid", fgColor="E2F0D9")
open_fill = PatternFill("solid", fgColor="FCE4D6")
resolved_fill = PatternFill("solid", fgColor="E2F0D9")
high_fill = PatternFill("solid", fgColor="F4CCCC")
thin = Side(style="thin", color="B7B7B7")
border = Border(left=thin, right=thin, top=thin, bottom=thin)

ws.merge_cells("A1:AA1")
ws["A1"] = "Medication Reconciliation — CT-2026-0713-18 — Mara Ellison"
ws["A1"].fill = title_fill
ws["A1"].font = Font(color="FFFFFF", bold=True, size=14)
ws["A1"].alignment = Alignment(horizontal="center")

ws.merge_cells("A2:AA2")
ws["A2"] = (
    "One row per retained medication identity. Source labels are crosswalked on that identity row; "
    "only active pharmacy-profile rows receive an operational action. Evidence and escalation are row-specific."
)
ws["A2"].alignment = Alignment(wrap_text=True)
ws["A2"].fill = subheader_fill

headers = [
    "Identity ID",
    "Retained medication identity",
    "Crosswalked source labels / identity basis",
    "Risk",
    "Escalation status",
    "Assigned owner",
    "Due date",
    "Source-supported disposition",
    "Evidence status",
    "Active profile row?",
    "Same-day profile action",
    "Profile action basis",
    "Evidence timeline",
    "Final outpatient directions",
    "Gap / next action",
    "Signed discharge summary",
    "Post-summary clarification",
    "Reconciliation queue",
    "Admission history",
    "Patient portal",
    "Pharmacy profile",
    "Dispense-history summary",
    "Allergy / safety",
    "Reviewer notes",
    "Verified disposition",
    "Verification owner",
    "Verified date",
]
for col, header in enumerate(headers, 1):
    cell = ws.cell(3, col, header)
    cell.fill = header_fill
    cell.font = Font(bold=True)
    cell.border = border
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

identity_rule = {
    "basal_insulin": "FG-SUB-001: approved U-100 glargine product identity",
    "ppi": "FG-SUB-002: approved pantoprazole-to-omeprazole outpatient crosswalk",
    "albuterol": "FG-SUB-003: approved albuterol HFA product identity",
    "potassium": "FG-SUB-004: approved potassium chloride ER 20 mEq product identity",
}

for identity_number, (key, display, _) in enumerate(MEDICATIONS, 1):
    row_number = identity_number + 3
    dispense_rows = by_dispense.get(key, [])
    if dispense_rows:
        latest = max(dispense_rows, key=lambda row: str(row["fill_date"]))
        dispense_summary = compact(
            [
                f"{len(dispense_rows)} row(s)",
                f"latest={latest['fill_date']}",
                f"status={latest['fill_status']}",
                f"qty={latest['quantity']}",
                f"days={latest['days_supply']}",
            ]
        )
    else:
        dispense_summary = ""

    source_labels = []
    for row in admission:
        if med_key(row["medication_label"]) == key:
            source_labels.append(row["medication_label"])
    for row in portal:
        if med_key(row["medication_label"]) == key:
            source_labels.append(row["medication_label"])
    for row in profile:
        if med_key(row["medication_label"]) == key:
            source_labels.append(row["medication_label"])
    for row in summary_rows:
        if med_key(row["Medication"]) == key:
            source_labels.append(row["Medication"])
    source_labels = list(dict.fromkeys(source_labels))
    identity_basis = "; ".join(source_labels)
    if key in identity_rule:
        identity_basis = compact([identity_basis, identity_rule[key]])

    active_profile_rows = [
        row
        for row in profile
        if med_key(row["medication_label"]) == key and row["profile_status"] == "active"
    ]
    has_active_profile = bool(active_profile_rows)
    fill_call = proposed[key] if has_active_profile else ""
    if fill_call == "KEEP":
        fill_call = "RELEASE"

    if not has_active_profile:
        if by_profile.get(key):
            fill_basis = (
                "No action: pharmacy-profile row is inactive. Retain the source disposition "
                "for crosswalk/history; do not infer a treatment or fill decision."
            )
        else:
            fill_basis = (
                "No action: no pharmacy-profile row exists for this identity. Retain the source "
                "disposition for crosswalk/history; route any signed order through normal intake."
            )
    elif fill_call == "RELEASE":
        fill_basis = "Active profile row plus explicit KEEP in the final signed summary."
    elif fill_call == "CHANGE":
        fill_basis = "Active profile row plus explicit CHANGE in the final signed summary."
    elif fill_call == "STOP":
        fill_basis = "Active profile row plus explicit STOP in the final signed summary."
    else:
        fill_basis = (
            "Active profile row conflicts with omission from the final summary and lacks a later "
            "resolution; send back for clarification."
        )

    if has_active_profile and key == "potassium":
        fill_basis = "RELEASE supported by the 2026-07-14 prescriber clarification after the reported BMP."
    elif has_active_profile and key == "ppi":
        fill_call = "RELEASE"
        fill_basis = (
            "RELEASE the existing omeprazole 20 mg profile identity: FG-SUB-002 resolves "
            "pantoprazole to omeprazole, and the 2026-07-14 prescriber clarification approves it."
        )
    elif has_active_profile and key == "lisinopril":
        fill_basis = "STOP supported by the final signed stop order and active HIGH angioedema allergy."
    elif has_active_profile and key == "spironolactone":
        fill_basis = "CLARIFY: final summary is silent; prescriber requested a cardiology decision."
    elif has_active_profile and key == "magnesium":
        fill_basis = "CLARIFY: final summary is silent; discharging-service answer is still requested."

    risk_tier = risk.get(key, "LOW")
    if fill_call == "CLARIFY":
        escalation_status = "OPEN ESCALATION"
        if key == "spironolactone":
            assigned_owner = "Owen Laird, PharmD — route to Cardiology"
        elif key == "magnesium":
            assigned_owner = "Owen Laird, PharmD — route to Discharging Service"
        else:
            assigned_owner = "Owen Laird, PharmD — obtain prescriber clarification"
    elif fill_call == "STOP":
        escalation_status = "ACTIVE PROFILE — STOP ACTION"
        assigned_owner = "Owen Laird, PharmD — Outpatient Pharmacy"
    elif fill_call == "CHANGE":
        escalation_status = "ACTIVE PROFILE — CHANGE ACTION"
        assigned_owner = "Owen Laird, PharmD — Outpatient Pharmacy"
    elif risk_tier == "HIGH":
        escalation_status = "HIGH-RISK REVIEW"
        assigned_owner = "Owen Laird, PharmD — Outpatient Pharmacy"
    elif risk_tier == "MED":
        escalation_status = "MED-RISK REVIEW"
        assigned_owner = "Owen Laird, PharmD — Outpatient Pharmacy"
    else:
        escalation_status = "NO OPEN ESCALATION"
        assigned_owner = "Owen Laird, PharmD — Outpatient Pharmacy"

    if fill_call or risk_tier == "HIGH":
        due_date = datetime(2026, 7, 15)
    elif risk_tier == "MED":
        due_date = datetime(2026, 7, 16)
    else:
        due_date = datetime(2026, 7, 17)

    evidence_events = []
    for row in admission:
        if med_key(row["medication_label"]) == key:
            evidence_events.append(
                (
                    str(row["captured_at"]),
                    f"{row['captured_at']} — Admission {row['source_row_id']} ({row['list_version']}): "
                    f"{row['medication_label']}; {row['reported_status']}; {row['sig']}",
                )
            )
    for row in portal:
        if med_key(row["medication_label"]) == key:
            evidence_events.append(
                (
                    str(row["entered_at"]),
                    f"{row['entered_at']} — Portal {row['source_row_id']}: {row['medication_label']}; "
                    f"{row['patient_status']}; {row['patient_comment']}",
                )
            )
    if dispense_rows:
        evidence_events.append(
            (
                str(latest["fill_date"]),
                f"{latest['fill_date']} — Latest dispense {latest['dispense_id']}: "
                f"{latest['medication_label']}; {latest['fill_status']}; "
                f"{latest['quantity']} units / {latest['days_supply']} days",
            )
        )
    for row in summary_rows:
        if med_key(row["Medication"]) == key:
            evidence_events.append(
                (
                    "2026-07-13 16:40",
                    f"2026-07-13 16:40 — Final summary {row['Line ID']}: "
                    f"{row['Disposition']} {row['Medication']}; {row['Dose / directions']}",
                )
            )
    for row in queue:
        if med_key(row["medication_or_scope"]) == key:
            evidence_events.append(
                (
                    row["event_time"],
                    f"{row['event_time']} — Queue {row['queue_id']}: {row['proposed_action']}; "
                    f"{row['queue_state']}/{row['revision_status']}",
                )
            )
    if key in clarifications:
        evidence_events.append(
            (
                "2026-07-14 15:10",
                f"2026-07-14 15:10 — Prescriber clarification: {clarifications[key].split(' — ', 1)[-1]}",
            )
        )
    related_allergies = []
    if key == "lisinopril":
        related_allergies = [row for row in allergies if row["substance"] == "lisinopril"]
    elif key == "amox_clav":
        related_allergies = [row for row in allergies if row["crosswalk_group"] == "penicillin"]
    elif key == "naproxen":
        related_allergies = [row for row in allergies if row["substance"] == "naproxen"]
    for row in related_allergies:
        evidence_events.append(
            (
                row["verified_at"],
                f"{row['verified_at']} — Allergy {row['allergy_id']}: {row['substance']} / "
                f"{row['reaction']}; {row['severity']} {row['status']}",
            )
        )
    for row in profile:
        if med_key(row["medication_label"]) == key:
            snapshot_time = source_exported_at["pharmacy_profile_snapshot.xlsx"]
            evidence_events.append(
                (
                    snapshot_time,
                    f"{snapshot_time} — Pharmacy snapshot {row['source_row_id']}: "
                    f"{row['medication_label']}; profile={row['profile_status']}; "
                    f"fill={row['fill_status']}",
                )
            )
    evidence_timeline = "\n".join(
        text
        for _, text in sorted(
            dict.fromkeys(evidence_events), key=lambda item: (item[0], item[1])
        )
    )

    values = [
        f"MED-{identity_number:03d}",
        display,
        identity_basis,
        risk_tier,
        escalation_status,
        assigned_owner,
        due_date,
        proposed[key],
        status[key],
        "YES" if has_active_profile else "NO",
        fill_call,
        fill_basis,
        evidence_timeline,
        summary_directions.get(key, ""),
        gap.get(key, ""),
        "\n".join(by_summary.get(key, [])),
        clarifications.get(key, ""),
        "\n".join(by_queue.get(key, [])),
        "\n".join(by_admission.get(key, [])),
        "\n".join(by_portal.get(key, [])),
        "\n".join(by_profile.get(key, [])),
        dispense_summary,
        allergy_notes.get(key, ""),
        "",
        "",
        "",
        "",
    ]
    for col, value in enumerate(values, 1):
        cell = ws.cell(row_number, col, value)
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    if status[key] == "OPEN":
        for col in range(1, len(headers) + 1):
            ws.cell(row_number, col).fill = open_fill
    elif status[key] == "RESOLVED":
        ws.cell(row_number, 9).fill = resolved_fill
    if risk.get(key) == "HIGH":
        ws.cell(row_number, 4).fill = high_fill

ws.freeze_panes = "H4"
ws.auto_filter.ref = f"A3:AA{ws.max_row}"
ws.row_dimensions[1].height = 24
ws.row_dimensions[2].height = 34
for row in range(4, ws.max_row + 1):
    ws.row_dimensions[row].height = 92

widths = {
    "A": 13,
    "B": 34,
    "C": 58,
    "D": 10,
    "E": 28,
    "F": 46,
    "G": 16,
    "H": 20,
    "I": 13,
    "J": 18,
    "K": 24,
    "L": 58,
    "M": 88,
    "N": 38,
    "O": 48,
    "P": 58,
    "Q": 48,
    "R": 50,
    "S": 58,
    "T": 52,
    "U": 58,
    "V": 27,
    "W": 42,
    "X": 30,
    "Y": 22,
    "Z": 22,
    "AA": 16,
}
for col, width in widths.items():
    ws.column_dimensions[col].width = width

profile_action_validation = DataValidation(
    type="list", formula1='"RELEASE,CHANGE,STOP,CLARIFY"', allow_blank=True
)
ws.add_data_validation(profile_action_validation)
profile_action_validation.add(f"K4:K{ws.max_row}")
disposition_validation = DataValidation(
    type="list", formula1='"KEEP,CHANGE,STOP,CLARIFY"', allow_blank=True
)
ws.add_data_validation(disposition_validation)
disposition_validation.add(f"Y4:Y{ws.max_row}")
date_validation = DataValidation(
    type="date",
    operator="between",
    formula1="DATE(2026,7,10)",
    formula2="DATE(2030,12,31)",
    allow_blank=True,
)
ws.add_data_validation(date_validation)
date_validation.add(f"G4:G{ws.max_row}")
date_validation.add(f"AA4:AA{ws.max_row}")
for row in range(4, ws.max_row + 1):
    ws.cell(row, 7).number_format = "yyyy-mm-dd"
    ws.cell(row, 27).number_format = "yyyy-mm-dd"

ws.conditional_formatting.add(
    f"A4:AA{ws.max_row}",
    FormulaRule(formula=["$I4=\"OPEN\""], fill=open_fill),
)

med_table = Table(displayName="MedRecTable", ref=f"A3:AA{ws.max_row}")
med_table.tableStyleInfo = TableStyleInfo(
    name="TableStyleMedium2",
    showFirstColumn=False,
    showLastColumn=False,
    showRowStripes=True,
    showColumnStripes=False,
)
ws.add_table(med_table)


summary_ws = wb.create_sheet("Summary_Checks", 1)
summary_ws.merge_cells("A1:D1")
summary_ws["A1"] = "Live Medication Reconciliation Summary And Checks"
summary_ws["A1"].fill = title_fill
summary_ws["A1"].font = Font(color="FFFFFF", bold=True, size=14)
summary_ws["A1"].alignment = Alignment(horizontal="center")
summary_ws.merge_cells("A2:D2")
summary_ws["A2"] = (
    "Formula-driven from MedRecTable. Edit Med_Reconciliation; counts and checks "
    "recalculate automatically when the workbook opens."
)
summary_ws["A2"].fill = subheader_fill
summary_ws["A2"].alignment = Alignment(wrap_text=True)

for col, value in enumerate(["Metric / check", "Live result", "Expected", "Status"], 1):
    cell = summary_ws.cell(4, col, value)
    cell.fill = header_fill
    cell.font = Font(bold=True)
    cell.border = border
    cell.alignment = Alignment(horizontal="center", wrap_text=True)

summary_metrics = [
    ("Retained medication identities", '=COUNTA(MedRecTable[Identity ID])', "35 retained rows", ""),
    ("Active profile identities", '=COUNTIF(MedRecTable[Active profile row?],"YES")', "Live count", ""),
    ("RELEASE actions", '=COUNTIF(MedRecTable[Same-day profile action],"RELEASE")', "Live count", ""),
    ("CHANGE actions", '=COUNTIF(MedRecTable[Same-day profile action],"CHANGE")', "Live count", ""),
    ("STOP actions", '=COUNTIF(MedRecTable[Same-day profile action],"STOP")', "Live count", ""),
    ("CLARIFY actions", '=COUNTIF(MedRecTable[Same-day profile action],"CLARIFY")', "Live count", ""),
    ("All active-profile actions", '=COUNTIF(MedRecTable[Same-day profile action],"<>")', "Live count", ""),
    ("HIGH-risk identities", '=COUNTIF(MedRecTable[Risk],"HIGH")', "Live count", ""),
    ("MED-risk identities", '=COUNTIF(MedRecTable[Risk],"MED")', "Live count", ""),
    ("LOW-risk identities", '=COUNTIF(MedRecTable[Risk],"LOW")', "Live count", ""),
    ("OPEN evidence statuses", '=COUNTIF(MedRecTable[Evidence status],"OPEN")', "Live count", ""),
    ("Due 2026-07-15", '=COUNTIF(MedRecTable[Due date],DATE(2026,7,15))', "Live count", ""),
    ("Due 2026-07-16", '=COUNTIF(MedRecTable[Due date],DATE(2026,7,16))', "Live count", ""),
    ("Due 2026-07-17", '=COUNTIF(MedRecTable[Due date],DATE(2026,7,17))', "Live count", ""),
]
for metric, formula, expected, status_text in summary_metrics:
    summary_ws.append([metric, formula, expected, status_text])

check_start = summary_ws.max_row + 2
for col, value in enumerate(["Integrity check", "Live exceptions", "Expected", "Status"], 1):
    cell = summary_ws.cell(check_start, col, value)
    cell.fill = header_fill
    cell.font = Font(bold=True)
    cell.border = border
    cell.alignment = Alignment(horizontal="center", wrap_text=True)

integrity_checks = [
    (
        "Active profile rows missing an action",
        '=COUNTIFS(MedRecTable[Active profile row?],"YES",MedRecTable[Same-day profile action],"")',
    ),
    (
        "Non-active rows carrying a profile action",
        '=COUNTIFS(MedRecTable[Active profile row?],"NO",MedRecTable[Same-day profile action],"<>")',
    ),
    (
        "Action count differs from active-profile count",
        '=ABS(COUNTIF(MedRecTable[Active profile row?],"YES")-COUNTIF(MedRecTable[Same-day profile action],"<>"))',
    ),
    (
        "Profile actions outside RELEASE/CHANGE/STOP/CLARIFY",
        '=COUNTIF(MedRecTable[Same-day profile action],"<>")'
        '-COUNTIF(MedRecTable[Same-day profile action],"RELEASE")'
        '-COUNTIF(MedRecTable[Same-day profile action],"CHANGE")'
        '-COUNTIF(MedRecTable[Same-day profile action],"STOP")'
        '-COUNTIF(MedRecTable[Same-day profile action],"CLARIFY")',
    ),
    ("Rows missing assigned owner", '=COUNTBLANK(MedRecTable[Assigned owner])'),
    ("Rows missing due date", '=COUNTBLANK(MedRecTable[Due date])'),
    ("Rows missing evidence timeline", '=COUNTBLANK(MedRecTable[Evidence timeline])'),
    (
        "CLARIFY action count differs from OPEN evidence count",
        '=ABS(COUNTIF(MedRecTable[Same-day profile action],"CLARIFY")-COUNTIF(MedRecTable[Evidence status],"OPEN"))',
    ),
    (
        "Duplicate retained identity IDs",
        '=SUMPRODUCT(--(COUNTIF(MedRecTable[Identity ID],MedRecTable[Identity ID])>1))',
    ),
]
for check, formula in integrity_checks:
    row_number = summary_ws.max_row + 1
    summary_ws.cell(row_number, 1, check)
    summary_ws.cell(row_number, 2, formula)
    summary_ws.cell(row_number, 3, 0)
    summary_ws.cell(row_number, 4, f'=IF(B{row_number}=C{row_number},"PASS","CHECK")')

for row in summary_ws.iter_rows(min_row=4, max_row=summary_ws.max_row, min_col=1, max_col=4):
    for cell in row:
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
for row in range(check_start + 1, summary_ws.max_row + 1):
    summary_ws.conditional_formatting.add(
        f"D{row}",
        FormulaRule(formula=[f'$D{row}="PASS"'], fill=resolved_fill),
    )
    summary_ws.conditional_formatting.add(
        f"D{row}",
        FormulaRule(formula=[f'$D{row}="CHECK"'], fill=open_fill),
    )
summary_ws.column_dimensions["A"].width = 54
summary_ws.column_dimensions["B"].width = 18
summary_ws.column_dimensions["C"].width = 20
summary_ws.column_dimensions["D"].width = 14
summary_ws.freeze_panes = "A5"


action_ws = wb.create_sheet("Action_View", 2)
action_ws.merge_cells("A1:L1")
action_ws["A1"] = "Recomputable Action View"
action_ws["A1"].fill = title_fill
action_ws["A1"].font = Font(color="FFFFFF", bold=True, size=14)
action_ws["A1"].alignment = Alignment(horizontal="center")
action_ws.merge_cells("A2:L2")
action_ws["A2"] = (
    "Linked to Med_Reconciliation. Use the table filters here for RELEASE, CHANGE, STOP, "
    "CLARIFY, owner, due date, risk, or evidence status; edit the source sheet, not this view."
)
action_ws["A2"].fill = subheader_fill
action_ws["A2"].alignment = Alignment(wrap_text=True)
view_headers = [
    "Identity ID",
    "Medication identity",
    "Risk",
    "Escalation status",
    "Assigned owner",
    "Due date",
    "Evidence status",
    "Active profile?",
    "Profile action",
    "Action basis",
    "Evidence timeline",
    "Gap / next action",
]
for col, header in enumerate(view_headers, 1):
    cell = action_ws.cell(4, col, header)
    cell.fill = header_fill
    cell.font = Font(bold=True)
    cell.border = border
    cell.alignment = Alignment(horizontal="center", wrap_text=True)

source_columns = ["A", "B", "D", "E", "F", "G", "I", "J", "K", "L", "M", "O"]
for index, source_row in enumerate(range(4, ws.max_row + 1), 5):
    for col, source_col in enumerate(source_columns, 1):
        cell = action_ws.cell(index, col, f"='Med_Reconciliation'!{source_col}{source_row}")
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    action_ws.cell(index, 6).number_format = "yyyy-mm-dd"
    action_ws.row_dimensions[index].height = 82

action_table = Table(displayName="ActionViewTable", ref=f"A4:L{action_ws.max_row}")
action_table.tableStyleInfo = TableStyleInfo(
    name="TableStyleMedium4",
    showFirstColumn=False,
    showLastColumn=False,
    showRowStripes=True,
    showColumnStripes=False,
)
action_ws.add_table(action_table)
action_ws.freeze_panes = "C5"
for col, width in enumerate([13, 34, 10, 28, 46, 16, 15, 16, 18, 58, 88, 48], 1):
    action_ws.column_dimensions[get_column_letter(col)].width = width
for row in range(5, action_ws.max_row + 1):
    for action, fill in [
        ("RELEASE", resolved_fill),
        ("CHANGE", PatternFill("solid", fgColor="FFF2CC")),
        ("STOP", high_fill),
        ("CLARIFY", open_fill),
    ]:
        action_ws.conditional_formatting.add(
            f"I{row}",
            FormulaRule(formula=[f'$I{row}="{action}"'], fill=fill),
        )


timeline = wb.create_sheet("Discharge_Timeline")
timeline_headers = ["Date/time", "Event", "Record", "Status / significance"]
for col, header in enumerate(timeline_headers, 1):
    cell = timeline.cell(1, col, header)
    cell.fill = header_fill
    cell.font = Font(bold=True)
    cell.border = border
timeline_rows = [
    ("2026-07-10 08:12", "Initial admission medication history captured", "admission_history_snapshot.xlsx / ADM-V1", "Historical home list"),
    ("2026-07-11 17:45", "Nurse-reviewed admission rows captured", "admission_history_snapshot.xlsx / ADM-V2", "Historical home list"),
    ("2026-07-12 06:05", "Admission history exported", "ADMEXP-20260712-4419", "Source snapshot"),
    ("2026-07-12 06:12", "Admission history review queue closed", "RCQ-7701", "Approved for pharmacist review"),
    ("2026-07-12 14:22", "Lisinopril angioedema allergy verified", "ALG-ACEI-20260712", "Active HIGH allergy"),
    ("2026-07-13 09:15", "Patient portal list captured", "PTLEXP-20260713-0915", "Patient-reported history"),
    ("2026-07-13 12:20", "Draft discharge summary issued", "DCS-744201-DRAFT", "Superseded"),
    ("2026-07-13 16:40", "Final medication summary signed/effective", "DCS-744201-FINAL v2.0", "Controls discharge regimen"),
    ("2026-07-13 16:44", "Final medication summary exported", "transition_summary.docx", "Final source snapshot"),
    ("2026-07-14 08:21", "Spironolactone clarification opened", "RCQ-7732", "OPEN"),
    ("2026-07-14 08:24", "Magnesium clarification opened", "RCQ-7737", "OPEN"),
    ("2026-07-14 11:42", "Partial prescriber response", "clarification email quoted reply", "Potassium held pending BMP"),
    ("2026-07-14 14:36", "BMP reportedly resulted: potassium 4.2", "clarification email only", "Underlying lab result absent from packet"),
    ("2026-07-14 15:10", "Prescriber clarification sent", "clarification_thread_july.eml", "Potassium, metoprolol and PPI resolved; spironolactone/magnesium unresolved"),
    ("2026-07-14 15:25", "Potassium queue item closed/current", "RCQ-7744", "Continue"),
    ("2026-07-14 17:05", "PPI queue item closed/current", "RCQ-7751", "Approved substitution"),
    ("2026-07-14 21:55", "Allergy extract exported", "allergy_status_extract.csv", "Latest allergy snapshot"),
    ("2026-07-14 22:05", "Reconciliation queue exported", "reconciliation_status_extract.csv", "Contains unresolved and stale preliminary rows"),
    ("2026-07-14 22:20", "Pharmacy profile exported", "pharmacy_profile_snapshot.xlsx", "Latest profile snapshot; not fully aligned"),
    ("2026-07-15", "Packet reconciliation date", "transition_packet_manifest.json", "No distinct discharge-completion event supplied"),
]
for row in timeline_rows:
    timeline.append(row)
for row in timeline.iter_rows():
    for cell in row:
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
timeline.freeze_panes = "A2"
timeline.auto_filter.ref = f"A1:D{timeline.max_row}"
for col, width in enumerate([22, 48, 38, 58], 1):
    timeline.column_dimensions[get_column_letter(col)].width = width


allergy_ws = wb.create_sheet("Allergies")
allergy_headers = list(allergies[0].keys())
allergy_ws.append(allergy_headers)
for cell in allergy_ws[1]:
    cell.fill = header_fill
    cell.font = Font(bold=True)
for row in allergies:
    allergy_ws.append([row[h] for h in allergy_headers])
for row in allergy_ws.iter_rows():
    for cell in row:
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
allergy_ws.freeze_panes = "A2"
allergy_ws.auto_filter.ref = allergy_ws.dimensions
for col in range(1, allergy_ws.max_column + 1):
    allergy_ws.column_dimensions[get_column_letter(col)].width = min(
        42,
        max(
            12,
            max(len(str(allergy_ws.cell(row, col).value or "")) for row in range(1, allergy_ws.max_row + 1)) + 2,
        ),
    )


source_ws = wb.create_sheet("Source_Index")
source_headers = ["Priority", "Source", "Exported at", "Purpose / operational use"]
source_ws.append(source_headers)
priority = {
    "transition_summary.docx": 1,
    "clarification_thread_july.eml": 2,
    "allergy_status_extract.csv": 3,
    "formulary_reference.docx": 4,
    "pharmacy_profile_snapshot.xlsx": 5,
    "admission_history_snapshot.xlsx": 6,
    "patient_list_snapshot.xlsx": 6,
    "reconciliation_status_extract.csv": 7,
}
for source in sorted(
    manifest["source_register"], key=lambda item: priority[item["relative_path"]]
):
    source_ws.append(
        [
            priority[source["relative_path"]],
            source["relative_path"],
            source["exported_at"],
            source["purpose"],
        ]
    )
for row in source_ws.iter_rows():
    for cell in row:
        cell.border = border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
for cell in source_ws[1]:
    cell.fill = header_fill
    cell.font = Font(bold=True)
source_ws.freeze_panes = "A2"
source_ws.auto_filter.ref = source_ws.dimensions
for col, width in enumerate([10, 38, 22, 68], 1):
    source_ws.column_dimensions[get_column_letter(col)].width = width


notes = wb.create_sheet("Read_Me")
notes.column_dimensions["A"].width = 125
readme = [
    "Purpose",
    "This workbook aligns medication evidence for case CT-2026-0713-18. It is an editable reconciliation aid, not a replacement for prescriber verification.",
    "",
    "How to use",
    "Edit Med_Reconciliation. Filter it by Evidence status=OPEN, Active profile row?=YES, action, owner, due date, or risk. Each MED-### row is one retained identity. Summary_Checks and Action_View are linked outputs that recalculate from the source table.",
    "",
    "Source rule",
    "The final signed care-transition summary controls after 2026-07-13 16:40. A later prescriber clarification can resolve a question created by the summary. Pharmacy, admission, and patient lists provide evidence but do not override those records.",
    "",
    "Normalization",
    "Approved product-level substitutions are grouped in one canonical row: glargine U-100 products, albuterol HFA products, potassium ER products, and pantoprazole/omeprazole. Explicit non-substitutions—metoprolol tartrate vs succinate ER and atorvastatin vs rosuvastatin—remain separate.",
    "",
    "Known open items",
    "Spironolactone and magnesium oxide are not confirmed. The actual BMP report, final dispensing proof, patient acknowledgment, and a distinct physical discharge timestamp are absent.",
    "",
    "Risk and escalation",
    "Each identity row contains a dated evidence timeline, controlled-reference risk tier, assigned owner, and exact due date. Dates are anchored to the packet reconciliation date of 2026-07-15: active-profile work and HIGH review are due 2026-07-15, MED follow-up 2026-07-16, and routine LOW reconciliation 2026-07-17. CLARIFY rows remain held after the due date until the named service responds.",
    "",
    "Owner assignment",
    "Owen Laird, PharmD is assigned as the accountable workbook owner because the source email identifies him as the outpatient pharmacist reconciling this case for discharge fill and the current queue rows name him as reviewer. Spironolactone routes to Cardiology and magnesium routes to the Discharging Service; Owen retains responsibility for the hold and follow-up.",
    "",
    "Recomputable checks and views",
    "Med_Reconciliation is an Excel table named MedRecTable. Summary_Checks uses COUNTIF/COUNTIFS formulas for live totals and zero-exception integrity checks. Action_View mirrors the row-level action fields with table filters. Workbook calculation is forced to automatic on open.",
    "",
    "Same-day fill call",
    "Only a row marked Active profile row?=YES receives a same-day profile action: RELEASE, CHANGE, STOP, or CLARIFY. Inactive or absent profile identities have no operational action. Their signed-summary disposition remains visible only as source evidence and is not converted into an inferred fill approval.",
]
for row_number, value in enumerate(readme, 1):
    notes.cell(row_number, 1, value)
    notes.cell(row_number, 1).alignment = Alignment(wrap_text=True, vertical="top")
    if value in {"Purpose", "How to use", "Source rule", "Normalization", "Known open items", "Risk and escalation", "Owner assignment", "Recomputable checks and views", "Same-day fill call"}:
        notes.cell(row_number, 1).font = Font(bold=True, color="FFFFFF")
        notes.cell(row_number, 1).fill = title_fill


for sheet in wb.worksheets:
    sheet.sheet_view.showGridLines = False

wb.calculation.calcMode = "auto"
wb.calculation.fullCalcOnLoad = True
wb.calculation.forceFullCalc = True
wb.calculation.calcOnSave = True
wb.save(OUTPUT)
print(f"Created {OUTPUT} with {len(MEDICATIONS)} medication rows")
